"""
validate_output.py

Pre-render validation gate. Runs after the document has been assembled
but before it is handed back, and checks it against
project_rulebook.json's global_constraints plus the structural
invariants the pipeline depends on.

Every finding names the paragraph index or section that triggered it.
Nothing here fails silently: checks that genuinely cannot be decided
without a layout engine (page overflow, blank pages) say so in their
finding rather than quietly passing.

Severity:
  "error"  — a stated rule is definitely violated.
  "warn"   — likely violated, or violated by a measure this can only
             approximate without laying out pages.
  "info"   — worth surfacing to the user but not a rule breach.
"""

import re

import rulebook as rb_module

# A TOC / List-of-X entry: trailing page number, usually with dot leaders.
LISTING_ENTRY_RE = re.compile(r"(\.{2,}\s*\d+\s*$)|(\s\d{1,3}\s*$)")
CHAPTER_HEADING_RE = re.compile(r"^CHAPTER\s+([A-Z]+|\d+)\b", re.IGNORECASE)

_DOT_LEADER_RE = re.compile(r"\.{2,}\s*\d+\s*$")
_TRAILING_PAGE_RE = re.compile(r"\s+\d{1,3}\s*$")


def _looks_like_listing_row(text):
    """True when the text has the shape of a contents-page row.

    Dot leaders followed by a number are unambiguous. A bare trailing
    number is not: "CHAPTER 2" ends in a number and is a perfectly good
    chapter heading. So a trailing number only counts when what precedes
    it is itself a title — two or more words — which "CHAPTER 2" is not
    once the number is removed, but "CHAPTER 2: Literature Review 8" is.
    """
    text = (text or "").strip()
    if not text:
        return False
    if _DOT_LEADER_RE.search(text):
        return True
    stripped = _TRAILING_PAGE_RE.sub("", text)
    if stripped == text:
        return False
    return len(stripped.split()) >= 2


def _finding(check, severity, message, where=None):
    return {"check": check, "severity": severity, "message": message, "where": where}


# ---------------------------------------------------------------------------
# Structural invariants (checked against the classification array)
# ---------------------------------------------------------------------------
def check_every_index_classified_once(classifications, paragraph_count):
    """Every input paragraph must appear exactly once. A missing index
    means a paragraph silently kept its original formatting; a duplicate
    means two roles fought over the same paragraph."""
    findings = []
    seen = {}
    for entry in classifications:
        idx = entry.get("index")
        seen[idx] = seen.get(idx, 0) + 1

    for idx, count in sorted(seen.items()):
        if count > 1:
            findings.append(_finding(
                "index_classified_once", "error",
                "Paragraph %d was classified %d times (roles must be exclusive)." % (idx, count),
                "paragraph %d" % idx))

    if paragraph_count is not None:
        missing = [i for i in range(paragraph_count) if i not in seen]
        for idx in missing:
            findings.append(_finding(
                "index_classified_once", "error",
                "Paragraph %d was never classified — it would be rendered with "
                "whatever formatting it already had." % idx,
                "paragraph %d" % idx))
        extra = [i for i in seen if i is None or i < 0 or i >= paragraph_count]
        for idx in extra:
            findings.append(_finding(
                "index_classified_once", "error",
                "Classification refers to paragraph index %r, which is outside "
                "the %d paragraphs that were extracted." % (idx, paragraph_count),
                "paragraph %r" % idx))
    return findings


def check_chapter_heading_zone_transitions(classifications, paragraph_texts):
    """A chapter_heading must genuinely open a chapter.

    The failure this catches is real and specific: an old Table of
    Contents contains lines like "CHAPTER ONE: INTRODUCTION ....... 1",
    which match a chapter-heading pattern. If one is tagged
    chapter_heading it both invents a chapter that doesn't exist and
    closes the stale-listing block early, so the rest of the old TOC
    leaks into the body as real content.
    """
    findings = []
    role_by_index = {c["index"]: c["role"] for c in classifications}
    order = sorted(role_by_index)

    def text_at(idx):
        return ((paragraph_texts[idx] if idx < len(paragraph_texts) else "") or "").strip()

    chapter_indices = [i for i in order if role_by_index[i] == "chapter_heading"]

    # The zone transition itself: once the first real chapter has opened,
    # the front matter is behind us, so no contents-page row may still be
    # arriving. Rows appearing after that point are exactly the leak this
    # check exists to catch.
    if chapter_indices:
        first_chapter = chapter_indices[0]
        for idx in order:
            if idx > first_chapter and role_by_index[idx] == "stale_listing_entry":
                findings.append(_finding(
                    "chapter_heading_zone", "error",
                    "Paragraph %d is a stale contents-page row but sits after the "
                    "first chapter heading (paragraph %d): %r. Either the chapter "
                    "heading is really a listing entry, or old listing rows are "
                    "leaking into the body."
                    % (idx, first_chapter, text_at(idx)[:70]),
                    "paragraph %d" % idx))

    # A real chapter heading is never followed immediately by a contents
    # row. When it is, the "chapter" is a line from the old table of
    # contents and the classifier has invented a chapter that doesn't
    # exist in the document.
    for pos, idx in enumerate(order):
        if role_by_index[idx] != "chapter_heading":
            continue
        nxt = order[pos + 1] if pos + 1 < len(order) else None
        if nxt is not None and role_by_index[nxt] == "stale_listing_entry":
            findings.append(_finding(
                "chapter_heading_zone", "error",
                "Paragraph %d is tagged chapter_heading but the next paragraph "
                "(%d) is a contents-page row: %r followed by %r. A chapter that "
                "opens straight into a listing is a listing entry, not a chapter."
                % (idx, nxt, text_at(idx)[:50], text_at(nxt)[:50]),
                "paragraph %d" % idx))
        elif _looks_like_listing_row(text_at(idx)):
            findings.append(_finding(
                "chapter_heading_zone", "warn",
                "Paragraph %d is tagged chapter_heading but has the shape of a "
                "contents-page entry (title followed by a page number): %r"
                % (idx, text_at(idx)[:70]),
                "paragraph %d" % idx))

    return findings


def check_no_stale_listing_in_output(doc, classifications, paragraph_texts):
    """stale_listing_entry paragraphs are old contents-page rows. They are
    deleted and the lists regenerated, so none may survive into the
    rendered document."""
    findings = []
    stale_texts = {
        (paragraph_texts[c["index"]] or "").strip()
        for c in classifications
        if c["role"] == "stale_listing_entry" and c["index"] < len(paragraph_texts)
    }
    stale_texts.discard("")
    if not stale_texts:
        return findings

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and text in stale_texts:
            findings.append(_finding(
                "stale_listing_leak", "error",
                "A paragraph classified stale_listing_entry survived into the "
                "final document at output position %d: %r" % (i, text[:70]),
                "output paragraph %d" % i))
    return findings


# ---------------------------------------------------------------------------
# Rulebook global_constraints
# ---------------------------------------------------------------------------
def check_no_blank_pages(doc):
    """Approximates the 'no blank pages' constraint.

    True page occupancy needs a layout engine, so this checks the thing
    that actually produces blank pages in practice: a page break with no
    renderable content between it and the next break.
    """
    from docx.oxml.ns import qn

    findings = []
    run_of_empties = 0
    last_break_at = None

    for i, para in enumerate(doc.paragraphs):
        has_text = bool(para.text.strip())
        xml = para._p.xml
        has_object = ("<w:drawing" in xml) or ("<w:pict" in xml)
        starts_page = bool(para.paragraph_format.page_break_before)
        has_manual_break = any(
            br.get(qn("w:type")) == "page"
            for run in para.runs for br in run._r.findall(qn("w:br"))
        )

        if starts_page or has_manual_break:
            if last_break_at is not None and run_of_empties >= 0 and not has_text and not has_object:
                findings.append(_finding(
                    "no_blank_pages", "warn",
                    "Two page breaks in a row with no content between them "
                    "(output paragraphs %d to %d) — this renders as a blank page."
                    % (last_break_at, i),
                    "output paragraphs %d-%d" % (last_break_at, i)))
            last_break_at = i
            run_of_empties = 0
            continue

        if has_text or has_object:
            run_of_empties = 0
            last_break_at = None
        else:
            run_of_empties += 1

    return findings


def _estimate_lines(text, chars_per_line=90):
    if not text:
        return 1
    return max(1, -(-len(text) // chars_per_line))


def check_single_page_sections(doc, max_lines=34):
    """The rulebook requires the cover page and certification page never
    to overflow onto a second page.

    Without a layout engine this can only be estimated, so it reports at
    "warn" and says so — a wrong guess here should prompt a human to look,
    not silently block the render.
    """
    findings = []
    # Which sections must fit on one page is the rulebook's call, not this
    # function's — it reads the no_overflow flag rather than carrying its
    # own list. Cover and title pages are built as fixed-height layout
    # tables and cannot overflow, so only the prose sections are watched.
    watched = {
        (section.get("title") or "").upper()
        for section in (rb_module.load_rulebook().get("sections") or [])
        if section.get("no_overflow") and section.get("title")
    }
    if not watched:
        return findings
    current = None
    lines = 0
    start_at = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        upper = text.upper()

        if upper in watched:
            current, lines, start_at = upper, 1, i
            continue

        if current:
            style = para.style.name if para.style else ""
            if para.paragraph_format.page_break_before or style == "Heading 1":
                if lines > max_lines:
                    findings.append(_finding(
                        "no_overflow", "warn",
                        "The %s section is about %d lines, which likely overflows "
                        "onto a second page (rulebook requires it to fit on one). "
                        "This is an estimate — page height isn't known without "
                        "laying the document out." % (current.title(), lines),
                        "output paragraphs %d-%d" % (start_at, i)))
                current = None
                continue
            lines += _estimate_lines(text)

    return findings


def check_abstract_word_count(doc, constraints):
    """Abstract must fall within the rulebook's word-count band."""
    findings = []
    lo = constraints.get("word_count_min")
    hi = constraints.get("word_count_max")
    if lo is None and hi is None:
        return findings

    collecting = False
    words = 0
    start_at = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.upper() == "ABSTRACT":
            collecting, words, start_at = True, 0, i
            continue
        if collecting:
            style = para.style.name if para.style else ""
            if style == "Heading 1" or para.paragraph_format.page_break_before:
                break
            words += len(text.split())

    if start_at is None:
        return findings  # no abstract in this document — not a violation here

    if lo is not None and words < lo:
        findings.append(_finding(
            "abstract_word_count", "warn",
            "Abstract is %d words; the rulebook requires %d-%d. The tool cannot "
            "lengthen it — the student needs to expand it." % (words, lo, hi),
            "abstract (output paragraph %d)" % start_at))
    elif hi is not None and words > hi:
        findings.append(_finding(
            "abstract_word_count", "warn",
            "Abstract is %d words; the rulebook allows %d-%d. The tool cannot "
            "shorten it without changing meaning — the student needs to trim it."
            % (words, lo, hi),
            "abstract (output paragraph %d)" % start_at))
    return findings


def check_references_sorted(doc):
    """References must be alphabetical by author surname."""
    findings = []
    collecting = False
    entries = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.upper() in ("REFERENCES", "BIBLIOGRAPHY"):
            collecting = True
            continue
        if collecting:
            style = para.style.name if para.style else ""
            if style == "Heading 1" or text.upper().startswith("APPENDI"):
                break
            if text:
                entries.append(text)

    if len(entries) < 2:
        return findings

    keys = [e.lower() for e in entries]
    for i in range(1, len(keys)):
        if keys[i] < keys[i - 1]:
            findings.append(_finding(
                "references_sorted", "error",
                "Reference list is not alphabetical: %r comes after %r."
                % (entries[i][:60], entries[i - 1][:60]),
                "reference entry %d" % i))
            break  # one report is enough; the whole list needs re-sorting
    return findings


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def check_abstract_keywords(doc, rb):
    """The rulebook requires the abstract to end with a keywords line."""
    findings = []
    spec = (rb_module.section_by_id("abstract", rb) or {}).get("keywords_line") or {}
    if not spec.get("required"):
        return findings
    prefix = (spec.get("prefix") or "Keywords:").lower()

    in_abstract = False
    found = False
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if text.upper() == "ABSTRACT":
            in_abstract = True
            continue
        if in_abstract:
            if style == "Heading 1" and text:
                break
            if text.lower().startswith(prefix):
                found = True
                break

    if in_abstract and not found:
        findings.append(_finding(
            "abstract_keywords", "warn",
            "The abstract has no %r line. The rulebook requires one, and the "
            "keywords have to come from the student — nothing here can "
            "invent them." % spec.get("prefix", "Keywords:"),
            "abstract"))
    return findings


def check_front_matter_order(doc, rb):
    """Front-matter sections must appear in the rulebook's declared order."""
    findings = []
    expected = [
        (s.get("title") or "").upper()
        for s in (rb.get("sections") or [])
        if s.get("title") and s.get("id") not in ("chapters", "references", "appendix")
    ]
    if not expected:
        return findings

    seen = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        if style != "Heading 1":
            continue
        title = para.text.strip().split("\n")[0].strip().upper()
        if title in expected and title not in seen:
            seen.append(title)

    wanted = [t for t in expected if t in seen]
    if seen != wanted:
        findings.append(_finding(
            "front_matter_order", "error",
            "Front matter is out of order. The rulebook asks for %s; the "
            "document has %s." % (" -> ".join(wanted), " -> ".join(seen)),
            "front matter"))
    return findings


def validate_document(doc, classifications, paragraph_texts, rb=None):
    """Runs every check and returns a structured report.

    Returns {"ok": bool, "errors": n, "warnings": n, "findings": [...],
             "rulebook_gaps": [...]}. `ok` is False only when a hard rule
    is definitely broken; approximations report as warnings so a wrong
    guess never blocks a render.
    """
    rb = rb or rb_module.load_rulebook()
    findings = []

    findings += check_every_index_classified_once(classifications, len(paragraph_texts))
    findings += check_chapter_heading_zone_transitions(classifications, paragraph_texts)
    findings += check_no_stale_listing_in_output(doc, classifications, paragraph_texts)
    findings += check_no_blank_pages(doc)
    findings += check_single_page_sections(doc)
    findings += check_abstract_word_count(doc, rb_module.abstract_constraints(rb))
    findings += check_abstract_keywords(doc, rb)
    findings += check_references_sorted(doc)
    findings += check_front_matter_order(doc, rb)

    errors = sum(1 for f in findings if f["severity"] == "error")
    warnings = sum(1 for f in findings if f["severity"] == "warn")

    return {
        "ok": errors == 0,
        "errors": errors,
        "warnings": warnings,
        "findings": findings,
        # Formatting inputs the rulebook doesn't pin down, surfaced so the
        # remaining hardcoded values stay visible.
        "rulebook_gaps": rb_module.unspecified_settings(rb),
    }
