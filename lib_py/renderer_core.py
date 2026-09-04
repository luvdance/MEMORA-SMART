"""
renderer_core.py

Opens the ORIGINAL uploaded .docx and edits it in place based on the
classification array. Tables (<w:tbl>) and images (<w:drawing>) are never
touched — the classifier only ever sees top-level paragraphs (see
extractor_core.py / paragraph_indexing.py), so there is no code path that
could alter them.

Passes implemented, matching what was agreed:
  0. Global font/spacing: Times New Roman 12pt, double-spaced body text;
     References section gets single spacing + 6pt space-before instead.
  2. Section break at the prelim/Chapter-One boundary: roman numerals
     (i, ii, iii...) before, arabic (1, 2, 3...) restarting at Chapter One.
  3. Heading styling + renumbering: chapter/section/subsection/subsubsection
     headings get real Word Heading styles (so the native TOC field can find
     them) and their numbering is REGENERATED from position in the document,
     not trusted from whatever the student typed. Chapters start on a fresh
     page; headings are kept with their following text.
  4. Figure/Table/Plate captions renumbered per chapter (numbering only —
     repositioning a caption relative to its image/table is intentionally
     NOT done yet; flagged as a follow-up, see README).
  5. References: duplicates removed, remaining entries alphabetized, hanging
     indent applied.
  6. Table of Contents inserted as a live native Word field (scans Heading
     1-3). List of Tables / List of Figures / List of Plates generated ONLY
     when that content type was actually found — never an empty list for a
     document that has no plates, etc. Page numbers in these three lists are
     STATIC for now (not live PAGEREF fields) — flagged as a follow-up.

Every structural section this produces is derived from what the classifier
actually found in THIS document — nothing is invented.
"""

import io
import re
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from paragraph_indexing import iter_indexed_paragraphs
import rulebook

FONT = "Times New Roman"  # module-level defaults, overridden per-call by options.formatting
BODY_PT = 12
CHAPTER_HEADING_PT = 14
SECTION_HEADING_PT = 12

CHAPTER_WORDS = ["ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE", "TEN"]

_ALIGNMENT_BY_NAME = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centered": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

CAPTION_PREFIX_RE = re.compile(
    r"^(fig(?:ure)?|table|plate)\.?\s*\d+(?:[.\-]\d+)*\s*[:.\-]?\s*", re.IGNORECASE
)
HEADING_NUMBER_PREFIX_RE = re.compile(r"^\d+(\.\d+)*\.?\s*")


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------
def strip_caption_prefix(text):
    return CAPTION_PREFIX_RE.sub("", text).strip()


def get_or_create_style(doc, name, style_type):
    """Returns the named style, creating it from scratch if the document
    doesn't already define it. Real-world documents — especially ones
    exported from Google Docs, stripped-down templates, or heavily "cleaned"
    files — often don't define Heading 1-4 at all, which used to crash the
    renderer outright. This is exactly the kind of poorly-formatted document
    the tool needs to handle, not an edge case to avoid."""
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def set_outline_level(paragraph, level):
    """Sets <w:outlineLvl w:val="level"/> directly on the paragraph (0-based:
    0=Heading 1, 1=Heading 2...). python-docx has no high-level API for this.
    Done explicitly on every heading paragraph — not just relied upon via the
    style — so the native Word TOC field (which scans by outline level) works
    correctly even when a freshly-created style doesn't carry the outline
    level Word's own built-in Heading styles normally carry automatically."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:outlineLvl"))
    if existing is not None:
        pPr.remove(existing)
    outline_el = OxmlElement("w:outlineLvl")
    outline_el.set(qn("w:val"), str(level))
    pPr.append(outline_el)


def apply_heading_style(doc, paragraph, level):
    """Applies both the named Heading style AND an explicit outline level —
    belt and suspenders, so the TOC field works regardless of whether the
    style itself behaves the way Word's real built-in Heading styles do."""
    style_name = f"Heading {level}"
    style = get_or_create_style(doc, style_name, WD_STYLE_TYPE.PARAGRAPH)
    paragraph.style = style
    set_outline_level(paragraph, level - 1)


def strip_heading_number_prefix(text):
    return HEADING_NUMBER_PREFIX_RE.sub("", text).strip()


def set_paragraph_text_single_run(paragraph, text, bold=None, size_pt=BODY_PT):
    """Replaces a paragraph's content with one run containing `text`,
    forcing font/size/bold explicitly (direct run formatting always beats
    style-derived formatting in Word, so this is the only reliable way to
    guarantee consistency regardless of whatever the student had)."""
    runs = paragraph.runs
    if runs:
        first = runs[0]
        first.text = text
        for extra in runs[1:]:
            extra.text = ""
        target_run = first
    else:
        target_run = paragraph.add_run(text)

    target_run.font.name = FONT
    target_run.font.size = Pt(size_pt)
    if bold is not None:
        target_run.font.bold = bold
    return target_run


def _apply_hanging_indent(paragraph, left_inches=None, hanging_inches=None,
                          tab_inches=None):
    """Applies the rulebook's indent geometry to one paragraph.

    A hanging indent is expressed as a positive number of inches in the
    rulebook but as a NEGATIVE first-line indent in Word, so the sign flip
    happens here, once, rather than at every call site.
    """
    pf = paragraph.paragraph_format
    if left_inches is not None:
        pf.left_indent = Inches(float(left_inches))
    if hanging_inches:
        hanging = float(hanging_inches)
        # The text body sits at left + hanging; the first line pulls back
        # to left, which is where the number prints.
        pf.left_indent = Inches(float(left_inches or 0) + hanging)
        pf.first_line_indent = Inches(-hanging)
    if tab_inches is not None:
        try:
            pf.tab_stops.add_tab_stop(Inches(float(tab_inches)))
        except Exception:
            pass  # a duplicate stop is harmless; never fail a render for one
    return paragraph


def _tabify_numbered_label(label, separator="tab"):
    """Turns "1.1 Background to the Study" into "1.1\tBackground to the
    Study" when the rulebook asks for a tab between the number and the
    words. A label with no leading number is returned untouched."""
    if separator != "tab" or not label:
        return label
    match = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)$", label)
    if not match:
        return label
    return match.group(1) + "\t" + match.group(2)


def set_two_line_heading(paragraph, line_1, line_2, bold=True, size_pt=None):
    """Writes a heading that reads as two centred lines but is ONE
    paragraph, the two halves separated by a soft line break.

    The rulebook requires the chapter number and the chapter title on
    separate lines. Two real paragraphs would give Word's table-of-contents
    field two separate entries for one chapter, so the break is a <w:br/>
    inside a single Heading 1 paragraph instead: the page looks the way the
    rulebook asks, and the contents page still shows one entry per chapter.
    """
    size_pt = size_pt or BODY_PT
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    first = paragraph.add_run(line_1)
    first.font.name = FONT
    first.font.size = Pt(size_pt)
    first.font.bold = bold

    if line_2:
        first._r.append(OxmlElement("w:br"))
        second = paragraph.add_run(line_2)
        second.font.name = FONT
        second.font.size = Pt(size_pt)
        second.font.bold = bold
    return paragraph


def _build_listing_rows(doc, heading_paragraph, section_id, entries, rb=None):
    """Writes the body of a List of Tables / Figures / Plates page.

    The rulebook gives each of those pages a two-part column header
    ("Table" on the left, "Page" on the right, italic) and an entry
    geometry: a hanging indent so a long caption wraps under its own text,
    and a right tab stop where the page number sits. Returns every
    paragraph it created, in order, so the caller can register them with
    the front-matter assembler.
    """
    spec = rulebook.listing_spec(section_id, rb)
    header = spec.get("column_header") or {}
    entry_rules = spec.get("entries") or {}
    tab_inches = entry_rules.get("page_number_tab_inches")
    created = []
    anchor = heading_paragraph

    def new_row():
        nonlocal anchor
        anchor = insert_paragraph_after(anchor, doc)
        created.append(anchor)
        return anchor

    if header.get("left") or header.get("right"):
        row = new_row()
        if tab_inches is not None:
            row.paragraph_format.tab_stops.add_tab_stop(
                Inches(float(tab_inches)), WD_TAB_ALIGNMENT.RIGHT)
        text = str(header.get("left") or "")
        if header.get("right"):
            text += "\t" + str(header["right"])
        run = row.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        run.font.italic = bool(header.get("italic"))

    for label in entries:
        row = new_row()
        if tab_inches is not None:
            row.paragraph_format.tab_stops.add_tab_stop(
                Inches(float(tab_inches)), WD_TAB_ALIGNMENT.RIGHT)
        _apply_hanging_indent(
            row, left_inches=0,
            hanging_inches=entry_rules.get("hanging_indent_inches"))
        row.paragraph_format.line_spacing = rulebook.line_spacing_value(
            entry_rules.get("line_spacing"), default=1.0)
        if entry_rules.get("space_after_pt") is not None:
            row.paragraph_format.space_after = Pt(entry_rules["space_after_pt"])
        # A list page repeats its subject in the column header ("Table",
        # "Figure"), so the entries themselves carry only the number and
        # the caption: "3.1<tab>Resistor Colour Code", not "Table 3.1: ...".
        entry_text = re.sub(r"^(Table|Figure|Fig|Plate)\s+", "", label,
                            flags=re.IGNORECASE)
        entry_text = re.sub(r"^(\d+(?:\.\d+)*)\s*[:.–-]\s*", r"\1 ", entry_text)
        run = row.add_run(_tabify_numbered_label(entry_text))
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)

    return created


def _apply_prelim_heading(doc, paragraph, section_id, fallback_title=None, rb=None):
    """Styles one front- or back-matter heading entirely from its section's
    `heading` block in the rulebook.

    Every such heading in the document — Declaration, Certification,
    Dedication, Acknowledgement, Table of Contents, the three List-of
    pages, Abstract, References, Appendix — goes through here, so their
    look is one rulebook edit away rather than nine identical literals
    scattered through the render pass.
    """
    section = rulebook.section_by_id(section_id, rb) or {}
    spec = section.get("heading", {}) or {}

    title = section.get("title") or fallback_title or paragraph.text.strip()
    if str(spec.get("case", "upper")).lower() == "upper":
        title = title.upper()

    size_pt = rulebook.heading_size_pt("prelim_heading", BODY_PT, rb)
    set_paragraph_text_single_run(
        paragraph, title, bold=spec.get("bold", True), size_pt=size_pt)
    apply_heading_style(doc, paragraph, 1)
    paragraph.alignment = _ALIGNMENT_BY_NAME.get(
        str(spec.get("alignment") or "center").lower(), WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.page_break_before = bool(
        section.get("page_break_before", True))
    if spec.get("space_after_pt") is not None:
        paragraph.paragraph_format.space_after = Pt(spec["space_after_pt"])
    return paragraph


def _apply_section_body(paragraph, section_id, rb=None, default_alignment="justify"):
    """Applies a front-matter section's `body` block: alignment and line
    spacing for its prose."""
    spec = rulebook.section_body_spec(section_id, rb) or {}
    paragraph.alignment = _ALIGNMENT_BY_NAME.get(
        str(spec.get("alignment") or default_alignment).lower(),
        WD_ALIGN_PARAGRAPH.JUSTIFY)
    paragraph.paragraph_format.line_spacing = rulebook.line_spacing_value(
        spec.get("line_spacing"), default=2.0)
    return paragraph


def enforce_run_fonts(paragraph, size_pt=BODY_PT):
    """For body paragraphs we don't rewrite text, but we DO force every run's
    font back to Times New Roman/12pt, since direct run-level formatting
    (whatever font/size the student happened to apply) otherwise overrides
    whatever we set at the style level."""
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(size_pt)


def insert_paragraph_after(anchor_paragraph, doc):
    """Creates a new empty paragraph and moves it to sit immediately after
    anchor_paragraph. Returns the new Paragraph object (from python-docx),
    so it can itself be used as the next anchor for another insertion,
    preserving order across multiple sequential inserts."""
    new_p = doc.add_paragraph()  # appended at end of body...
    anchor_paragraph._p.addnext(new_p._p)  # ...then relocated right here
    return new_p


def set_page_numbering(sectPr_element, fmt, start):
    """Sets/creates <w:pgNumType w:fmt=".." w:start=".."/> on a sectPr,
    respecting OOXML's required child ordering (pgNumType must come after
    pgMar and before cols, if present)."""
    existing = sectPr_element.find(qn("w:pgNumType"))
    if existing is not None:
        sectPr_element.remove(existing)

    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))

    # Insert in the correct schema position: after pgMar if present,
    # otherwise after pgSz, otherwise just append (still valid — Word is
    # tolerant of pgNumType appearing later than pgMar in practice, but we
    # do this properly rather than relying on that tolerance).
    pgMar = sectPr_element.find(qn("w:pgMar"))
    if pgMar is not None:
        pgMar.addnext(pgNumType)
    else:
        sectPr_element.append(pgNumType)


# ---------------------------------------------------------------------------
# Table borders ("open"/academic style)
# ---------------------------------------------------------------------------
_TBLPR_TAIL_TAGS = (
    "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
    "w:tblCaption", "w:tblDescription",
)


def _insert_in_tblPr_order(tblPr, new_el):
    """tblPr's children have a required schema order; w:tblBorders must
    come before w:shd/w:tblLayout/w:tblCellMar/w:tblLook/etc if present.
    Insert before the first such tail element, or append if none exist."""
    for tag in _TBLPR_TAIL_TAGS:
        tail_el = tblPr.find(qn(tag))
        if tail_el is not None:
            tail_el.addprevious(new_el)
            return
    tblPr.append(new_el)


def _border_edge(tag, val, sz=4, color="000000"):
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), val)
    if val != "nil":
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
    return el


def apply_academic_table_borders(table, rb=None):
    """Academic 'open' table style: a single line above the header row
    (table top), a single line below the header row, a single line
    closing the bottom of the table, and NO other borders — no outer box
    sides, no vertical lines, no lines between body rows.

    Table-level tblBorders top/bottom only ever affect the outer boundary
    of the whole table (top of the first row, bottom of the last row) —
    NOT every row — so setting top=single/bottom=single there already
    gives two of the three required lines for free. insideH/insideV/left/
    right are all set to "nil" so no other line appears anywhere. The
    third line (below the header row specifically) needs a per-cell
    tcBorders override on row 0 only, since insideH governs ALL inter-row
    lines uniformly and can't target just one row boundary.

    Any pre-existing per-cell tcBorders overrides (e.g. leftover from a
    "Table Grid"-style table, which is exactly the messy starting point
    real student documents have) are stripped first, so they can't
    reintroduce lines the table-wide setting just removed.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    existing_borders = tblPr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    # Which edges are drawn comes from table_rules.borders in the
    # rulebook. "none" becomes OOXML's "nil"; anything else is passed
    # through as the line style Word should draw.
    border_rules = (rulebook.table_rules(rb) or {}).get("borders") or {}

    def edge_val(key, default):
        value = border_rules.get(key, default)
        return "nil" if str(value).lower() in ("none", "nil", "") else str(value)

    borders = OxmlElement("w:tblBorders")
    for key, tag, default in (
        ("outer_top", "top", "single"),
        ("outer_bottom", "bottom", "single"),
        ("outer_left", "left", "none"),
        ("outer_right", "right", "none"),
        ("inside_horizontal", "insideH", "none"),
        ("inside_vertical", "insideV", "none"),
    ):
        borders.append(_border_edge(tag, edge_val(key, default)))
    _insert_in_tblPr_order(tblPr, borders)

    header_bottom = edge_val("header_row_bottom", "single")
    header_bold = bool((rulebook.table_rules(rb) or {}).get("header_row_bold"))
    cell_align = _ALIGNMENT_BY_NAME.get(
        str((rulebook.table_rules(rb) or {}).get("cell_alignment") or "").lower())

    rows = table.rows
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            existing_tc_borders = tcPr.find(qn("w:tcBorders"))
            if existing_tc_borders is not None:
                tcPr.remove(existing_tc_borders)

            if cell_align is not None:
                for cell_para in cell.paragraphs:
                    cell_para.alignment = cell_align

            if r_idx == 0:
                if header_bold:
                    for cell_para in cell.paragraphs:
                        for cell_run in cell_para.runs:
                            cell_run.font.bold = True
                # Header row: explicit bottom line. Every other edge is
                # left unset so it falls through to the table-wide
                # tblBorders above (which already gives the correct
                # top/nil-everything-else behavior).
                tc_borders = OxmlElement("w:tcBorders")
                tc_borders.append(_border_edge("bottom", header_bottom))
                # tcPr child order: tcBorders comes early (after tcW,
                # before shd/tcMar/etc) — insert at the front is safe
                # since tcW (if present) is python-docx-managed and this
                # mirrors the tblPr helper's approach at cell scope.
                first_child = tcPr.find(qn("w:tcW"))
                if first_child is not None:
                    first_child.addnext(tc_borders)
                else:
                    tcPr.insert(0, tc_borders)


def apply_academic_borders_to_all_tables(doc, rb=None):
    """Applies the open academic border style to every table in the
    document, including tables nested inside table cells (doc.tables only
    lists top-level tables)."""
    from docx.table import Table

    for tbl_element in doc.element.body.iter(qn("w:tbl")):
        apply_academic_table_borders(Table(tbl_element, doc), rb)


# ---------------------------------------------------------------------------
# Blank-page cleanup
# ---------------------------------------------------------------------------
def _paragraph_is_empty(paragraph):
    """True if a paragraph has no visible content at all — no text, and no
    image/object either (so an image-only caption-less paragraph, which
    legitimately has empty .text, is never mistaken for blank filler)."""
    if paragraph.text.strip():
        return False
    xml = paragraph._p.xml
    return "<w:drawing" not in xml and "<w:pict" not in xml and "<w:object" not in xml


def _element_is_image_paragraph(element):
    """True for a <w:p> whose content is a picture — the 'figure' a figure
    caption belongs to."""
    if element.tag != qn("w:p"):
        return False
    xml = element.xml
    return "<w:drawing" in xml or "<w:pict" in xml or "<w:object" in xml


def _element_is_skippable(element):
    """Filler between a caption and the thing it captions: an empty
    paragraph carrying neither text nor a picture."""
    if element.tag != qn("w:p"):
        return False
    if _element_is_image_paragraph(element):
        return False
    return not "".join(element.itertext()).strip()


def _nearest_sibling(element, matcher, forward, max_steps=3):
    """Walk siblings in one direction, stepping over blank filler, and
    return the first element the matcher accepts."""
    sibling = element.getnext() if forward else element.getprevious()
    steps = 0
    while sibling is not None and steps < max_steps:
        if matcher(sibling):
            return sibling
        if not _element_is_skippable(sibling):
            return None
        sibling = sibling.getnext() if forward else sibling.getprevious()
        steps += 1
    return None


def reposition_captions(table_captions, figure_captions, rb):
    """Put every caption on the side of its table/figure the rulebook asks
    for.

    project_rulebook.json states table_rules.caption_position and
    figure_rules.caption_position, and those two strings are the only
    input here — this function has no opinion of its own about which way
    round they go.

    Runs at the very end of render(), on the finished XML tree. Moving a
    paragraph changes what a positional index would refer to, so it has to
    happen after every index-dependent pass has finished, per this file's
    index-safety rule.
    """
    moved = 0

    def place(caption_paragraphs, target_matcher, want_above):
        nonlocal moved
        for para in caption_paragraphs:
            cap = para._p
            if cap.getparent() is None:
                continue  # deleted earlier in the render
            after = _nearest_sibling(cap, target_matcher, forward=True)
            before = _nearest_sibling(cap, target_matcher, forward=False)
            if want_above and after is None and before is not None:
                # Caption sits below its target but belongs above it.
                cap.getparent().remove(cap)
                before.addprevious(cap)
                moved += 1
            elif not want_above and before is None and after is not None:
                # Caption sits above its target but belongs below it.
                cap.getparent().remove(cap)
                after.addnext(cap)
                moved += 1

    table_pos = (rulebook.table_rules(rb) or {}).get("caption_position", "above")
    figure_pos = (rulebook.figure_rules(rb) or {}).get("caption_position", "below")

    place(table_captions, lambda el: el.tag == qn("w:tbl"),
          want_above=(str(table_pos).lower() == "above"))
    place(figure_captions, _element_is_image_paragraph,
          want_above=(str(figure_pos).lower() == "above"))
    return moved


def remove_manual_page_breaks_and_blank_pages(doc):
    """Real student documents very commonly use manual page breaks
    (Word's Ctrl+Enter, a <w:br w:type="page"/> run character) between
    prelim sections — Declaration, Certification, Table of Contents, List
    of Tables, etc. This renderer inserts its OWN page_break_before flag
    on exactly the headings that need one; a leftover manual break from
    the original file is redundant at best, and produces a genuine BLANK
    PAGE at worst — e.g. a manual break immediately followed by a heading
    that ALSO carries page_break_before doubles up, leaving an empty page
    between them. This is exactly what real-world testing found.

    Two passes:
      1. Strip every manual page-break run in the document. If that break
         was a paragraph's ONLY content, remove the now-empty paragraph
         entirely rather than leaving a stray blank line behind.
      2. Collapse any remaining fully-empty paragraph (no text, no image)
         that sits immediately before a paragraph carrying
         page_break_before=True — that heading's own break already starts
         a fresh page, so the blank spacer before it does nothing but add
         risk of an extra blank page.
    """
    for p in list(doc.paragraphs):
        changed = False
        for run in list(p.runs):
            for br in run._r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    run._r.remove(br)
                    changed = True
        if changed and _paragraph_is_empty(p):
            parent = p._p.getparent()
            if parent is not None:
                parent.remove(p._p)

    for p in list(doc.paragraphs):
        if not _paragraph_is_empty(p):
            continue
        p_el = p._p
        if p_el.getparent() is None:
            continue  # already removed above in this same sweep
        next_el = p_el.getnext()
        while next_el is not None and next_el.tag != qn("w:p"):
            next_el = next_el.getnext()
        if next_el is None:
            continue
        from docx.text.paragraph import Paragraph
        next_p = Paragraph(next_el, p._parent)
        if next_p.paragraph_format.page_break_before:
            parent = p_el.getparent()
            if parent is not None:
                parent.remove(p_el)


# ---------------------------------------------------------------------------
# Front-matter assembly
# ---------------------------------------------------------------------------
# The canonical order every Nigerian university expects its preliminary
# pages in. Everything before Chapter One is ASSEMBLED into this order —
# it is not left wherever it happened to sit in the student's file, and
# not left wherever the generator happened to insert it.
# Derived from project_rulebook.json's `sections` array — the order the
# document is assembled in is a JSON edit, not a code change, and the
# shipped order cannot drift from the spec. See rulebook.prelim_rank().
# Front-matter ordering, derived from the rulebook's own `sections` array.
# Module level is only a convenience for helpers called outside a render;
# render() recomputes it per call (see below) so that reordering sections
# in the JSON reorders the document without needing a process restart.
PRELIM_RANK = rulebook.prelim_rank()


class PrelimAssembler:
    """Collects every piece of front matter — whether PRESERVED from the
    student's original file or freshly GENERATED by this renderer — each
    tagged with its canonical rank, then lays them all out in canonical
    order immediately before Chapter One in a single pass.

    This exists because front matter is the one part of the document that
    genuinely has to be ASSEMBLED rather than edited in place. Everything
    from Chapter One onward is edited where it sits (which is what keeps
    tables and images byte-identical), but the preliminary pages have a
    fixed conventional order that real student files routinely violate,
    and that generated replacement sections have no way of slotting into
    on their own.

    The bug that motivated this: preserved content and generated content
    used two different, mutually unaware placement mechanisms — generated
    blocks anchored at the very start of the document, preserved blocks
    left where they were. A student whose file had no Acknowledgements
    got an AI-drafted one inserted as page i, ahead of their own cover
    page. Routing BOTH kinds through one ranked assembler is what makes
    that structurally impossible rather than merely fixed.
    """

    def __init__(self):
        self._items = []   # (rank, sequence, element)
        self._seq = 0

    def add(self, rank, obj):
        """Registers a paragraph, table, or raw element at `rank`.
        Sequence is preserved within a rank, so a block's internal line
        order survives the sort."""
        if obj is None:
            return
        el = getattr(obj, "_p", None)
        if el is None:
            el = getattr(obj, "_tbl", obj)
        self._items.append((rank, self._seq, el))
        self._seq += 1

    def add_all(self, rank, objs):
        for obj in objs:
            self.add(rank, obj)

    def layout(self, first_chapter_paragraph):
        """Detaches every registered element and re-inserts it in
        canonical order directly before Chapter One.

        Elements deleted earlier in the pipeline (e.g. an original
        Declaration whose toggle asked for a regenerated replacement)
        report a None parent and are dropped here rather than
        resurrected. Chapter One onward is never touched."""
        if first_chapter_paragraph is None:
            return
        live = [item for item in self._items if item[2].getparent() is not None]
        if not live:
            return
        live.sort(key=lambda item: (item[0], item[1]))

        anchor = first_chapter_paragraph._p
        for _, _, el in live:
            el.getparent().remove(el)
        for _, _, el in live:
            anchor.addprevious(el)


# ---------------------------------------------------------------------------
# Reference cleanup
# ---------------------------------------------------------------------------
def _normalize_for_dedupe(text):
    normalized = re.sub(r"[^\w\s]", "", text.lower())
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized[:60]


def clean_and_sort_references(raw_texts):
    seen = {}
    for text in raw_texts:
        key = _normalize_for_dedupe(text)
        if not key:
            continue
        existing = seen.get(key)
        if not existing or len(text) > len(existing):
            seen[key] = text.strip()
    return sorted(seen.values(), key=lambda s: s.lower())


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------
def render(file_bytes: bytes, classifications: list, options: dict = None):
    options = options or {}
    doc = Document(io.BytesIO(file_bytes))
    role_by_index = {c["index"]: c["role"] for c in classifications}

    # --- Apply formatting overrides for this call ---------------------------
    # Simplification: these helpers read FONT/BODY_PT/etc as module globals
    # rather than taking a config parameter each. Reassigning them here makes
    # a single render() call fully configurable without threading a config
    # object through every helper function. Trade-off: not safe for two
    # concurrent renders in the same warm process — acceptable for now since
    # each Vercel invocation completes before the next starts; would need a
    # proper refactor (pass config explicitly) if that ever changes.
    global FONT, BODY_PT, CHAPTER_HEADING_PT, SECTION_HEADING_PT, PRELIM_RANK
    # Every default here comes from project_rulebook.json — nothing is
    # invented in this file. The user's own formatting overrides (font,
    # size, spacing chosen in the UI) still win, but where they say
    # nothing the rulebook decides, not a hardcoded literal.
    rb = rulebook.load_rulebook()
    rb_defaults = rulebook.document_defaults(rb)
    # Recomputed per render so the order of `sections` in the JSON is what
    # decides the order of the finished document, every time.
    PRELIM_RANK = rulebook.prelim_rank(rb)

    fmt_opts = options.get("formatting", {})
    FONT = fmt_opts.get("fontFamily") or rb_defaults.get("font")
    BODY_PT = fmt_opts.get("fontSizePt") or rb_defaults.get("font_size_pt")
    CHAPTER_HEADING_PT = rulebook.heading_size_pt("chapter_heading", BODY_PT, rb)
    SECTION_HEADING_PT = rulebook.heading_size_pt("section_heading", BODY_PT, rb)
    PRELIM_HEADING_PT = rulebook.heading_size_pt("prelim_heading", BODY_PT, rb)

    # Layout rules for the two heading families the body uses. Read once,
    # here, so no branch below reaches for a literal of its own.
    _chapter_heading_rules = rulebook.chapter_heading_rules(rb)
    _section_heading_rules = rulebook.section_heading_rules(rb)
    _reference_rules = rulebook.reference_rules(rb)
    body_line_spacing = fmt_opts.get("lineSpacing") or rulebook.line_spacing_value(
        rb_defaults.get("line_spacing"))
    body_alignment = _ALIGNMENT_BY_NAME.get(
        (rb_defaults.get("body_alignment") or "justify").lower(),
        WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Page margins — specified by the rulebook and, until now, never
    # applied at all.
    margins = rb_defaults.get("margins_inches") or {}
    for section in doc.sections:
        if "top" in margins:
            section.top_margin = Inches(margins["top"])
        if "bottom" in margins:
            section.bottom_margin = Inches(margins["bottom"])
        if "left" in margins:
            section.left_margin = Inches(margins["left"])
        if "right" in margins:
            section.right_margin = Inches(margins["right"])

    prelim_toggles = options.get("prelimToggles", {})
    personal = options.get("personalDetails", {})
    ai_content = options.get("aiDraftedContent", {})
    directives = options.get("directives", {})

    def get_role(idx):
        return role_by_index.get(idx, "body_paragraph")

    # --- Configure base styles once, document-wide -------------------------
    normal = get_or_create_style(doc, "Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.line_spacing = body_line_spacing
    normal.paragraph_format.space_after = Pt(0)

    for level, pt_size in [(1, CHAPTER_HEADING_PT), (2, SECTION_HEADING_PT),
                            (3, SECTION_HEADING_PT), (4, SECTION_HEADING_PT)]:
        style = get_or_create_style(doc, f"Heading {level}", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = Pt(pt_size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        if level == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Table borders: academic "open" style on every table -----------------
    # Independent of paragraph indices (tables are never part of
    # iter_indexed_paragraphs), so safe to run at any point in render().
    apply_academic_borders_to_all_tables(doc, rb)

    # --- PASS A: compute all renumbering labels + collect list entries -----
    # (read-only pass — we need to know EVERY caption's final label before
    # we can build the List of Tables/Figures/Plates, which physically sits
    # earlier in the document than most of the captions themselves.)
    chapter_ctr = section_ctr = subsection_ctr = subsubsection_ctr = 0
    figure_ctr = table_ctr = plate_ctr = 0

    heading_labels = {}   # index -> new text
    chapter_title_labels = {}  # chapter anchor index -> that chapter's TITLE
    # line. The rulebook puts a chapter heading on two lines ("CHAPTER ONE"
    # then "INTRODUCTION"), so the number line and the title line are
    # tracked apart and joined by a line break at render time.
    heading_levels = {}   # index -> effective Heading level (2/3/4) — see
    # the orphan-heading promotion in the section/subsection/subsubsection
    # branches below; this can differ from the role's "natural" level.
    caption_labels = {}   # index -> new text

    # --- Front-matter registration -----------------------------------------
    # Every paragraph BEFORE Chapter One is registered with the assembler,
    # tagged with the canonical section it belongs to. Done here in PASS A
    # (read-only, so indices still exactly match the classification array)
    # because doing it later would be unsafe: paragraph deletions in PASS B
    # shift what a fresh index scan would number each remaining paragraph,
    # per this file's index-safety rule. The assembler lays everything out
    # in canonical order after PASS B and generate_prelim_pages have both
    # had their say — see PrelimAssembler for why generated and preserved
    # content must go through the same mechanism.
    PRELIM_HEADING_ROLE_RANK = {
        # Literal front-matter heading lines. deterministicClassify emits
        # these as first-class roles; the previous AI classifier lumped
        # them into prelim_label and they had to be recovered by matching
        # the heading text.
        "declaration_heading": PRELIM_RANK["declaration"],
        "certification_heading": PRELIM_RANK["certification"],
        "dedication_heading": PRELIM_RANK["dedication"],
        "acknowledgement_heading": PRELIM_RANK["acknowledgement"],
        "declaration_body": PRELIM_RANK["declaration"],
        "certification_body": PRELIM_RANK["certification"],
        "dedication_body": PRELIM_RANK["dedication"],
        "acknowledgement_body": PRELIM_RANK["acknowledgement"],
        "abstract_heading": PRELIM_RANK["abstract"],
        "abstract_body": PRELIM_RANK["abstract"],
        "toc_heading": PRELIM_RANK["toc"],
        "list_of_tables_heading": PRELIM_RANK["list_of_tables"],
        "list_of_figures_heading": PRELIM_RANK["list_of_figures"],
        "list_of_plates_heading": PRELIM_RANK["list_of_plates"],
    }

    # Where the cover/title page stops and the rest of the front matter
    # begins. Everything before the first real front-matter or body
    # heading is cover-page furniture: the title, "BY", the student's
    # name, the matric number, the submission sentence, the supervisor
    # line, the date.
    #
    # This matters because deterministicClassify tags all of those lines
    # prelim_label, and prelim_label's default treatment is "style this as
    # a section heading on its own page". Applied to cover lines that is
    # both wrong and duplicated — the regenerated cover page already
    # carries the same information — so when the cover is being rebuilt,
    # the originals are removed instead.
    #
    # Computed once, here, from the untouched classification array. PASS B
    # deletes paragraphs, so an index derived later would refer to a
    # different paragraph than the classifier meant.
    _ZONE_ENDING_ROLES = set(PRELIM_HEADING_ROLE_RANK) | {
        "chapter_heading", "references_heading", "appendix_heading",
        "section_heading", "subsection_heading", "subsubsection_heading",
        "body_paragraph",
    }
    cover_zone_end = None
    for _i in sorted(role_by_index):
        if role_by_index[_i] in _ZONE_ENDING_ROLES:
            cover_zone_end = _i
            break
    if cover_zone_end is None:
        cover_zone_end = -1  # no front matter at all — treat nothing as cover

    prelim = PrelimAssembler()
    current_prelim_rank = PRELIM_RANK["front"]
    toc_figure_entries = []
    toc_table_entries = []
    toc_plate_entries = []
    table_caption_paragraphs = []
    figure_caption_paragraphs = []
    reference_indices = []
    reference_texts = []
    first_chapter_index = None
    chapter_continuation_indices = set()  # paragraphs merged into the
    # previous chapter_heading rather than treated as a new chapter — see
    # prev_role handling below.
    prev_role = None
    current_chapter_anchor_idx = None

    for idx, p in iter_indexed_paragraphs(doc):
        role = get_role(idx)
        text = p.text.strip()

        if first_chapter_index is None and not (role == "chapter_heading" and prev_role != "chapter_heading"):
            # Still in the prelim area (haven't reached the real start of
            # Chapter One yet on this iteration). cover_title/prelim_label
            # are EXPLICIT front-matter signals from the classifier — treat
            # them as always resetting to the front bucket regardless of
            # where they physically sit, since a label like "SUPERVISOR:"
            # or the student's name sitting in the middle of the file
            # (e.g. physically after a misplaced Acknowledgements section)
            # is still unambiguously front matter, not a continuation of
            # whatever heading happens to precede it in the file.
            label_upper = text.upper()
            if role == "cover_title":
                current_prelim_rank = PRELIM_RANK["front"]
            elif role == "prelim_label":
                if "DECLARATION" in label_upper:
                    current_prelim_rank = PRELIM_RANK["declaration"]
                elif "CERTIF" in label_upper:
                    current_prelim_rank = PRELIM_RANK["certification"]
                elif "DEDICAT" in label_upper:
                    current_prelim_rank = PRELIM_RANK["dedication"]
                elif "ACKNOWLEDG" in label_upper:
                    current_prelim_rank = PRELIM_RANK["acknowledgement"]
                else:
                    current_prelim_rank = PRELIM_RANK["front"]
            elif role in PRELIM_HEADING_ROLE_RANK:
                current_prelim_rank = PRELIM_HEADING_ROLE_RANK[role]
            # else (body_paragraph, stale_listing_entry, table_data, a
            # stray caption, etc.): carries forward whatever section rank
            # it's physically sitting under — reasonable default, and this
            # is genuinely ambiguous content to attribute otherwise.
            prelim.add(current_prelim_rank, p)

        if role == "chapter_heading":
            if prev_role == "chapter_heading":
                # Continuation of the SAME chapter heading, not a new
                # chapter — very common real-world pattern: student put the
                # chapter number ("CHAPTER ONE") and the chapter title
                # ("INTRODUCTION") on two separate paragraphs. Merge into
                # the anchor paragraph's label instead of incrementing.
                extra = strip_heading_number_prefix(text)
                extra = re.sub(r"^CHAPTER\s+(\w+|\d+)\s*[:.\-]?\s*", "", extra, flags=re.IGNORECASE)
                if extra:
                    existing = chapter_title_labels.get(current_chapter_anchor_idx, "")
                    joiner = " " if existing else ""
                    chapter_title_labels[current_chapter_anchor_idx] = (
                        existing + joiner + extra.upper())
                chapter_continuation_indices.add(idx)
                prev_role = role
                continue

            chapter_ctr += 1
            section_ctr = subsection_ctr = subsubsection_ctr = 0
            figure_ctr = table_ctr = plate_ctr = 0
            if first_chapter_index is None:
                first_chapter_index = idx
            current_chapter_anchor_idx = idx
            word = CHAPTER_WORDS[chapter_ctr - 1] if chapter_ctr <= len(CHAPTER_WORDS) else str(chapter_ctr)
            title_part = strip_heading_number_prefix(text)
            title_part = re.sub(r"^CHAPTER\s+(\w+|\d+)\s*[:.\-]?\s*", "", title_part, flags=re.IGNORECASE)
            # The number line and the title line, kept apart. Which form the
            # document actually used does not matter: a student who wrote
            # "CHAPTER 1: Introduction" on one line and one who used two
            # paragraphs both arrive here with the same pair.
            heading_labels[idx] = rulebook.fill_template(
                (_chapter_heading_rules.get("line_1") or {}).get("template")
                or "CHAPTER {chapter_word}",
                {"chapter_word": word, "chapter_number": chapter_ctr})
            chapter_title_labels[idx] = title_part.upper() if title_part else ""
        elif role == "section_heading":
            section_ctr += 1
            subsection_ctr = subsubsection_ctr = 0
            heading_labels[idx] = f"{chapter_ctr}.{section_ctr} {strip_heading_number_prefix(text)}"
            heading_levels[idx] = 2
        elif role == "subsection_heading":
            if section_ctr == 0:
                # "Alien" heading: a subsection with no section ancestor
                # yet in this chapter — a real pattern found in testing
                # (the classifier tagged "2.5.1 Common Extraction Methods"
                # as a subsection right after "CHAPTER TWO", with no
                # "2.x" section heading between them). Numbering it
                # normally would produce a broken "2.0.1" label. It's
                # structurally the first grouping heading in the chapter,
                # so promote it to behave as a section instead — both the
                # numbering AND the actual Heading-level style/outline
                # level applied in PASS B (see heading_levels) reflect
                # the promotion, so the native TOC field nests it
                # correctly too.
                section_ctr += 1
                subsection_ctr = subsubsection_ctr = 0
                heading_labels[idx] = f"{chapter_ctr}.{section_ctr} {strip_heading_number_prefix(text)}"
                heading_levels[idx] = 2
            else:
                subsection_ctr += 1
                subsubsection_ctr = 0
                heading_labels[idx] = f"{chapter_ctr}.{section_ctr}.{subsection_ctr} {strip_heading_number_prefix(text)}"
                heading_levels[idx] = 3
        elif role == "subsubsection_heading":
            if section_ctr == 0:
                # Same orphan situation, two levels deep — promote all
                # the way to section level.
                section_ctr += 1
                subsection_ctr = subsubsection_ctr = 0
                heading_labels[idx] = f"{chapter_ctr}.{section_ctr} {strip_heading_number_prefix(text)}"
                heading_levels[idx] = 2
            elif subsection_ctr == 0:
                # Has a section ancestor but no subsection ancestor —
                # promote to subsection level.
                subsection_ctr += 1
                subsubsection_ctr = 0
                heading_labels[idx] = f"{chapter_ctr}.{section_ctr}.{subsection_ctr} {strip_heading_number_prefix(text)}"
                heading_levels[idx] = 3
            else:
                subsubsection_ctr += 1
                heading_labels[idx] = (
                    f"{chapter_ctr}.{section_ctr}.{subsection_ctr}.{subsubsection_ctr} "
                    f"{strip_heading_number_prefix(text)}"
                )
                heading_levels[idx] = 4
        elif role == "figure_caption":
            figure_ctr += 1
            label = f"Figure {chapter_ctr}.{figure_ctr}: {strip_caption_prefix(text)}"
            caption_labels[idx] = label
            toc_figure_entries.append(label)
        elif role == "table_caption":
            table_ctr += 1
            label = f"Table {chapter_ctr}.{table_ctr}: {strip_caption_prefix(text)}"
            caption_labels[idx] = label
            toc_table_entries.append(label)
        elif role == "plate_caption":
            plate_ctr += 1
            label = f"Plate {chapter_ctr}.{plate_ctr}: {strip_caption_prefix(text)}"
            caption_labels[idx] = label
            toc_plate_entries.append(label)
        elif role == "reference_entry":
            reference_indices.append(idx)
            reference_texts.append(text)

        prev_role = role

    cleaned_references = clean_and_sort_references(reference_texts)

    has_dedication = any(get_role(idx) == "dedication_body" for idx, _ in iter_indexed_paragraphs(doc))
    has_acknowledgement = any(get_role(idx) == "acknowledgement_body" for idx, _ in iter_indexed_paragraphs(doc))

    # --- PASS B: mutate the real document -----------------------------------
    first_chapter_paragraph = None

    HEADING_ROLES = {
        "chapter_heading", "section_heading", "subsection_heading", "subsubsection_heading",
        "references_heading", "appendix_heading", "toc_heading",
        "list_of_tables_heading", "list_of_figures_heading", "list_of_plates_heading",
        "abstract_heading", "prelim_label",
    }

    for idx, p in list(iter_indexed_paragraphs(doc)):
        role = get_role(idx)

        # Safety net: if this paragraph isn't one we're explicitly assigning
        # a Heading style to, but it already carries a stray Heading style
        # from the original (unreliable) student formatting, normalize it
        # back to Normal. Left alone, junk like a title-page name line
        # accidentally styled "Heading 2" would show up as a bogus entry in
        # the native Word TOC field, which scans by style, not content.
        if role not in HEADING_ROLES and p.style is not None and p.style.name.startswith("Heading"):
            p.style = get_or_create_style(doc, "Normal", WD_STYLE_TYPE.PARAGRAPH)

        if role == "chapter_heading" and idx in chapter_continuation_indices:
            # Merged into the previous chapter_heading's label already —
            # this physical paragraph is no longer needed.
            p._p.getparent().remove(p._p)

        elif role == "chapter_heading":
            _line1_rules = _chapter_heading_rules.get("line_1") or {}
            _line2_rules = _chapter_heading_rules.get("line_2") or {}
            apply_heading_style(doc, p, 1)
            if _chapter_heading_rules.get("layout") == "two_lines":
                set_two_line_heading(
                    p, heading_labels[idx], chapter_title_labels.get(idx, ""),
                    bold=_line1_rules.get("bold", True),
                    size_pt=CHAPTER_HEADING_PT)
            else:
                _joined = heading_labels[idx]
                _title = chapter_title_labels.get(idx, "")
                if _title:
                    _joined = _joined + ": " + _title
                set_paragraph_text_single_run(
                    p, _joined, bold=_line1_rules.get("bold", True),
                    size_pt=CHAPTER_HEADING_PT)
            p.paragraph_format.page_break_before = bool(
                _chapter_heading_rules.get("page_break_before", True))
            p.paragraph_format.keep_with_next = bool(
                _chapter_heading_rules.get("keep_with_next", True))
            p.paragraph_format.alignment = _ALIGNMENT_BY_NAME.get(
                str(_line2_rules.get("alignment")
                    or _line1_rules.get("alignment") or "center").lower(),
                WD_ALIGN_PARAGRAPH.CENTER)
            if _chapter_heading_rules.get("space_after_pt") is not None:
                p.paragraph_format.space_after = Pt(
                    _chapter_heading_rules["space_after_pt"])
            if first_chapter_paragraph is None:
                first_chapter_paragraph = p

        elif role in ("section_heading", "subsection_heading", "subsubsection_heading"):
            # heading_levels holds the EFFECTIVE level computed in PASS A,
            # which can differ from the role's natural level when an
            # orphan heading (no section/subsection ancestor yet in this
            # chapter) was promoted — keeps the Heading-style/outline
            # level consistent with the numbering actually printed, so
            # the native TOC field nests it correctly.
            level = heading_levels.get(
                idx, {"section_heading": 2, "subsection_heading": 3, "subsubsection_heading": 4}[role]
            )
            set_paragraph_text_single_run(
                p,
                _tabify_numbered_label(
                    heading_labels[idx],
                    _section_heading_rules.get("number_separator", "tab")),
                bold=_section_heading_rules.get("bold", True),
                size_pt=SECTION_HEADING_PT)
            p.style = get_or_create_style(doc, f"Heading {level}", WD_STYLE_TYPE.PARAGRAPH)
            set_outline_level(p, level - 1)
            p.paragraph_format.keep_with_next = bool(
                _section_heading_rules.get("keep_with_next", True))
            # The rulebook puts the number at the margin and the words at a
            # tab stop, so a heading that wraps lines up under its own text
            # rather than under its number.
            _apply_hanging_indent(
                p,
                left_inches=_section_heading_rules.get("left_indent_inches"),
                hanging_inches=_section_heading_rules.get("hanging_indent_inches"),
                tab_inches=_section_heading_rules.get("tab_stop_inches"))
            if _section_heading_rules.get("space_before_pt") is not None:
                p.paragraph_format.space_before = Pt(_section_heading_rules["space_before_pt"])
            if _section_heading_rules.get("space_after_pt") is not None:
                p.paragraph_format.space_after = Pt(_section_heading_rules["space_after_pt"])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif role in ("figure_caption", "table_caption", "plate_caption"):
            # Table captions and figure captions follow different rules —
            # the rulebook indents a table caption's title to a tab stop
            # and centres a figure caption under its image — so each reads
            # its own block rather than sharing one style.
            _cap_rules = (rulebook.table_rules(rb) if role == "table_caption"
                          else rulebook.figure_rules(rb)) or {}
            _cap_text = caption_labels[idx]
            if _cap_rules.get("caption_number_separator") == "tab":
                _cap_text = re.sub(
                    r"^(Table|Figure|Fig|Plate)\s+(\d+(?:\.\d+)*)\s*[:.]?\s*",
                    r"\1 \2\t", _cap_text, flags=re.IGNORECASE)
            set_paragraph_text_single_run(
                p, _cap_text, bold=_cap_rules.get("caption_bold", True))
            p.alignment = _ALIGNMENT_BY_NAME.get(
                str(_cap_rules.get("caption_alignment") or "center").lower(),
                WD_ALIGN_PARAGRAPH.CENTER)
            _apply_hanging_indent(
                p,
                left_inches=_cap_rules.get("caption_left_indent_inches"),
                hanging_inches=_cap_rules.get("caption_hanging_indent_inches"),
                tab_inches=_cap_rules.get("caption_tab_stop_inches"))
            if _cap_rules.get("caption_space_before_pt") is not None:
                p.paragraph_format.space_before = Pt(_cap_rules["caption_space_before_pt"])
            if _cap_rules.get("caption_space_after_pt") is not None:
                p.paragraph_format.space_after = Pt(_cap_rules["caption_space_after_pt"])
            # A caption above its target must stay with it; one below must
            # stay with what precedes it, so keep_with_next would be wrong.
            p.paragraph_format.keep_with_next = (
                str(_cap_rules.get("caption_position", "above")).lower() == "above")
            # Held for the caption-placement pass at the end of render().
            # Collected as live paragraph objects rather than indices so
            # the later pass never has to re-derive a position that
            # deletions here have already shifted.
            if role == "table_caption":
                table_caption_paragraphs.append(p)
            else:
                figure_caption_paragraphs.append(p)

        elif role == "reference_entry":
            p._p.getparent().remove(p._p)  # deleted here; cleaned list re-inserted below

        elif role == "stale_listing_entry":
            p._p.getparent().remove(p._p)  # old TOC/List-of-X entry, regenerated below instead

        elif role == "references_heading":
            _apply_prelim_heading(
                doc, p, "references",
                fallback_title=p.text.strip().upper() or "REFERENCES")
            anchor = p
            for ref_text in cleaned_references:
                anchor = insert_paragraph_after(anchor, doc)
                run = anchor.add_run(ref_text)
                run.font.name = FONT
                run.font.size = Pt(BODY_PT)
                _apply_hanging_indent(
                    anchor,
                    left_inches=0,
                    hanging_inches=_reference_rules.get("hanging_indent_inches"))
                anchor.paragraph_format.line_spacing = rulebook.line_spacing_value(
                    _reference_rules.get("line_spacing"), default=1.0)
                if _reference_rules.get("space_after_pt") is not None:
                    anchor.paragraph_format.space_after = Pt(
                        _reference_rules["space_after_pt"])
                anchor.alignment = _ALIGNMENT_BY_NAME.get(
                    str(_reference_rules.get("alignment") or "justify").lower(),
                    WD_ALIGN_PARAGRAPH.JUSTIFY)

        elif role == "appendix_heading":
            _apply_prelim_heading(doc, p, "appendix",
                                  fallback_title=p.text.strip().upper())

        elif role == "toc_heading":
            # Toggle defaults to True when the caller doesn't specify it,
            # so callers that don't pass prelimToggles at all (existing
            # tests, older callers) keep the original always-on behavior.
            if prelim_toggles.get("tableOfContents", True):
                _apply_prelim_heading(doc, p, "table_of_contents", fallback_title="TABLE OF CONTENTS")
                field_p = insert_paragraph_after(p, doc)
                _insert_toc_field(field_p)
                # The heading itself was registered in PASS A; this field
                # paragraph is brand new, so it has to be registered too or
                # the assembler would relocate the heading and strand the
                # field behind it.
                prelim.add(PRELIM_RANK["toc"], field_p)
            else:
                p._p.getparent().remove(p._p)

        elif role == "list_of_tables_heading":
            if toc_table_entries and prelim_toggles.get("listOfTables", True):
                _apply_prelim_heading(doc, p, "list_of_tables", fallback_title="LIST OF TABLES")
                _rows = _build_listing_rows(doc, p, "list_of_tables", toc_table_entries, rb)
                for _row in _rows:
                    # Newly created, so register alongside the heading.
                    prelim.add(PRELIM_RANK["list_of_tables"], _row)
            else:
                # No tables were actually found in this document, or the
                # user unchecked "List of Tables" — remove the stray
                # heading entirely rather than leave an empty/unwanted list.
                p._p.getparent().remove(p._p)

        elif role == "list_of_figures_heading":
            if toc_figure_entries and prelim_toggles.get("listOfFigures", True):
                _apply_prelim_heading(doc, p, "list_of_figures", fallback_title="LIST OF FIGURES")
                _rows = _build_listing_rows(doc, p, "list_of_figures", toc_figure_entries, rb)
                for _row in _rows:
                    # Newly created, so register alongside the heading.
                    prelim.add(PRELIM_RANK["list_of_figures"], _row)
            else:
                p._p.getparent().remove(p._p)

        elif role == "list_of_plates_heading":
            if toc_plate_entries and prelim_toggles.get("listOfPlates", True):
                _apply_prelim_heading(doc, p, "list_of_plates", fallback_title="LIST OF PLATES")
                _rows = _build_listing_rows(doc, p, "list_of_plates", toc_plate_entries, rb)
                for _row in _rows:
                    # Newly created, so register alongside the heading.
                    prelim.add(PRELIM_RANK["list_of_plates"], _row)
            else:
                p._p.getparent().remove(p._p)

        elif role in ("abstract_heading",):
            # Toggle off preserves the original abstract as ordinary
            # content rather than deleting it — it's substantive student
            # writing, not a derived listing, so "not wanted as its own
            # styled prelim page" must never mean "destroyed".
            if prelim_toggles.get("abstract", True):
                _apply_prelim_heading(doc, p, "abstract", fallback_title=p.text.strip())
            else:
                enforce_run_fonts(p)

        elif role == "cover_title" and (prelim_toggles.get("coverPage") or prelim_toggles.get("titlePage")):
            p._p.getparent().remove(p._p)  # replaced by fresh template below

        elif role == "declaration_body" and prelim_toggles.get("declaration"):
            p._p.getparent().remove(p._p)

        elif role == "certification_body" and prelim_toggles.get("certification"):
            p._p.getparent().remove(p._p)

        elif role in ("declaration_heading", "certification_heading",
                      "dedication_heading", "acknowledgement_heading"):
            # First-class front-matter heading roles from
            # deterministicClassify. Same treatment the prelim_label
            # branch below gives a heading it recognised by text, but
            # driven by the role itself rather than string matching.
            _toggle_for_role = {
                "declaration_heading": "declaration",
                "certification_heading": "certification",
            }.get(role)
            if _toggle_for_role and prelim_toggles.get(_toggle_for_role):
                # A regenerated replacement section is coming, so the
                # original heading goes with the original body.
                p._p.getparent().remove(p._p)
            else:
                _apply_prelim_heading(doc, p, {
                    "declaration_heading": "declaration",
                    "certification_heading": "certification",
                    "dedication_heading": "dedication",
                    "acknowledgement_heading": "acknowledgement",
                }[role], fallback_title=p.text.strip().upper())

        elif role == "prelim_label":
            label_upper = p.text.strip().upper()
            in_cover_zone = idx < cover_zone_end
            should_delete = (
                # Cover-page furniture, and a fresh cover page is being
                # built from the personal details — keeping these would
                # print the student's name, matric number and supervisor
                # twice, each on its own page.
                (in_cover_zone and (prelim_toggles.get("coverPage")
                                    or prelim_toggles.get("titlePage")))
                or ("DECLARATION" in label_upper and prelim_toggles.get("declaration"))
                or ("CERTIFICATION" in label_upper and prelim_toggles.get("certification"))
                # Dedication/Acknowledgement labels are NOT deleted here even
                # if toggled on — original personal writing is preserved
                # when present; AI drafting only fills in when it's ABSENT
                # (handled separately in generate_prelim_pages).
            )
            if should_delete:
                p._p.getparent().remove(p._p)
            elif in_cover_zone:
                # Cover lines the user chose to keep. They are lines on one
                # page, not section headings: centred, body-sized, no page
                # break between them, and deliberately not a Heading style
                # so they never appear as entries in the generated table of
                # contents.
                set_paragraph_text_single_run(p, p.text.strip().upper(), bold=True, size_pt=BODY_PT)
                p.style = get_or_create_style(doc, "Normal", WD_STYLE_TYPE.PARAGRAPH)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.page_break_before = False
            else:
                set_paragraph_text_single_run(p, p.text.strip().upper(), bold=True, size_pt=CHAPTER_HEADING_PT)
                apply_heading_style(doc, p, 1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.page_break_before = True

        elif role == "abstract_body":
            # Abstract formatting is pinned by the rulebook's own
            # constraints block (single spacing, its own point size,
            # justified), not by the user's body-text choices.
            _abs = rulebook.abstract_constraints(rb)
            enforce_run_fonts(p, size_pt=_abs.get("font_size_pt") or BODY_PT)
            p.paragraph_format.line_spacing = rulebook.line_spacing_value(
                _abs.get("line_spacing"), default=1.0)
            p.alignment = _ALIGNMENT_BY_NAME.get(
                (_abs.get("alignment") or "justify").lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)

        elif role in ("body_paragraph", "appendix_body"):
            # Flowing body prose. Force consistent double-spacing (or
            # whatever the user configured) and justification EXPLICITLY at
            # the paragraph level — not just via the Normal style — because
            # direct paragraph-level formatting (whatever line_spacing/
            # alignment the student's original paragraph happened to carry)
            # always wins over style-derived formatting in Word. Relying on
            # the Normal style alone left chapters that had their own
            # conflicting direct formatting (e.g. pasted in from another
            # document) inconsistent with the rest of the document — a real
            # bug found in testing.
            enforce_run_fonts(p)
            p.paragraph_format.line_spacing = body_line_spacing
            p.alignment = body_alignment  # rulebook document_defaults.body_alignment
            # Block paragraphs: the rulebook sets both of these to zero, so
            # a student's stray first-line indent or paragraph spacing is
            # normalised away rather than carried through.
            if rb_defaults.get("body_first_line_indent_inches") is not None:
                p.paragraph_format.first_line_indent = Inches(
                    rb_defaults["body_first_line_indent_inches"])
            if rb_defaults.get("body_space_after_pt") is not None:
                p.paragraph_format.space_after = Pt(rb_defaults["body_space_after_pt"])

        elif role in ("dedication_body", "acknowledgement_body"):
            # The student's own dedication and acknowledgement prose, kept
            # word for word but given the alignment and spacing their
            # rulebook section asks for. These are flowing paragraphs, not
            # hand-built layouts, so reflowing them is safe — unlike the
            # signature blocks handled by the catch-all below.
            enforce_run_fonts(p)
            _apply_section_body(
                p,
                "dedication" if role == "dedication_body" else "acknowledgement",
                rb)

        else:
            # declaration_body, certification_body, cover_title,
            # table_data, etc.
            # Content untouched — just force consistent font/size. This is
            # also what protects tab+underscore signature lines on the
            # Certification page: we never touch paragraph structure/tabs,
            # only the run-level font, so hand-built layouts survive intact.
            # These roles deliberately do NOT get the body_paragraph
            # spacing/justify treatment above — declaration/certification/
            # dedication/acknowledgement pages follow their own prelim-page
            # conventions (see generate_prelim_pages), and cover_title/
            # table_data layouts must not be reflowed.
            enforce_run_fonts(p)


    # A stable element to build brand-new front matter in front of. Its
    # exact position doesn't matter — the assembler repositions everything
    # at the end — it only has to be a real, attached element.
    _prelim_build_anchor = (
        first_chapter_paragraph._p if first_chapter_paragraph is not None
        else (doc.paragraphs[0]._p if doc.paragraphs else None)
    )

    # --- Generate replacement / AI-drafted prelim pages ---------------------
    # Must happen AFTER Pass B (which depends on stable original-document
    # indices). Everything it creates registers with the SAME assembler the
    # preserved originals registered with in PASS A, so generated and
    # preserved front matter get placed by one mechanism rather than two
    # mutually unaware ones.
    generate_prelim_pages(
        doc, prelim_toggles, personal, ai_content,
        has_dedication=has_dedication, has_acknowledgement=has_acknowledgement,
        prelim=prelim, before_element=_prelim_build_anchor,
    )

    # --- Fallback: build a TOC / List-of-X the original never had -----------
    # These are built anywhere convenient and registered at their canonical
    # rank; the assembler below is what actually positions them.
    if (prelim_toggles.get("tableOfContents", True)
            and not any(get_role(i) == "toc_heading" for i, _ in iter_indexed_paragraphs(doc))):
        # The document had NO Table of Contents at all. Unlike the
        # List-of-X fallbacks (gated on entries actually existing), a TOC
        # is always worth building fresh: chapter and section headings
        # always exist in a real project, and the native TOC field
        # populates itself from them the moment the file is opened (see
        # _enable_update_fields_on_open).
        prelim.add_all(PRELIM_RANK["toc"],
                       _build_toc_section(doc, _prelim_build_anchor, rb))
    if (toc_table_entries and prelim_toggles.get("listOfTables", True)
            and not any(get_role(i) == "list_of_tables_heading" for i, _ in iter_indexed_paragraphs(doc))):
        prelim.add_all(PRELIM_RANK["list_of_tables"],
                       _build_list_section(doc, _prelim_build_anchor, "list_of_tables",
                                           "LIST OF TABLES", toc_table_entries, rb))
    if (toc_figure_entries and prelim_toggles.get("listOfFigures", True)
            and not any(get_role(i) == "list_of_figures_heading" for i, _ in iter_indexed_paragraphs(doc))):
        prelim.add_all(PRELIM_RANK["list_of_figures"],
                       _build_list_section(doc, _prelim_build_anchor, "list_of_figures",
                                           "LIST OF FIGURES", toc_figure_entries, rb))
    if (toc_plate_entries and prelim_toggles.get("listOfPlates", True)
            and not any(get_role(i) == "list_of_plates_heading" for i, _ in iter_indexed_paragraphs(doc))):
        prelim.add_all(PRELIM_RANK["list_of_plates"],
                       _build_list_section(doc, _prelim_build_anchor, "list_of_plates",
                                           "LIST OF PLATES", toc_plate_entries, rb))

    # --- Assemble the front matter in canonical order -----------------------
    # Single placement pass over everything registered above — preserved
    # and generated alike. Runs before the blank-page cleanup so that pass
    # sees the final paragraph adjacency, and before Pass C so the section
    # break lands at the true prelim/Chapter-One boundary.
    prelim.layout(first_chapter_paragraph)

    # --- Blank-page cleanup ---------------------------------------------
    # Must run AFTER Pass B and generate_prelim_pages, since those are
    # what determine the final set of page_break_before headings this
    # checks against; must run BEFORE Pass C's section-break normalization
    # so that pass sees the final, cleaned-up paragraph sequence.
    remove_manual_page_breaks_and_blank_pages(doc)

    # --- PASS C: section break at the prelim / Chapter One boundary --------
    # First, normalize away any section breaks the ORIGINAL document already
    # had (common — e.g. many students already split off an unnumbered cover
    # page). Left alone, our new break would stack on top of theirs and
    # reset roman-numeral counting partway through the prelim pages
    # (i, ii, i, ii, iii...). Removing all pre-existing inline section
    # breaks merges the whole document back into one section first, so our
    # own single clean break produces continuous numbering regardless of
    # whatever structure the original file happened to have.
    # (Known trade-off: this also discards any page-specific settings from
    # those original breaks, e.g. a landscape orientation page for a wide
    # table — flagged as a follow-up if this turns out to matter in practice.)
    for p_element in list(doc.element.body.iterchildren()):
        if p_element.tag == qn("w:p"):
            pPr = p_element.find(qn("w:pPr"))
            if pPr is not None:
                inline_sectPr = pPr.find(qn("w:sectPr"))
                if inline_sectPr is not None:
                    pPr.remove(inline_sectPr)

    if first_chapter_paragraph is not None:
        prev_p_element = first_chapter_paragraph._p.getprevious()
        # Walk back to the nearest actual <w:p> sibling (skip over anything
        # else, e.g. a table, that might sit immediately before — rare but
        # safer to handle).
        while prev_p_element is not None and prev_p_element.tag != qn("w:p"):
            prev_p_element = prev_p_element.getprevious()

        final_sectPr = doc.sections[-1]._sectPr

        if prev_p_element is not None:
            # There IS preliminary content before Chapter One — split it
            # into its own roman-numbered section.
            prelim_sectPr = deepcopy(final_sectPr)
            pPr = prev_p_element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                prev_p_element.insert(0, pPr)
            pPr.append(prelim_sectPr)
            set_page_numbering(prelim_sectPr, "lowerRoman", 1)
            _ensure_page_number_footer(doc.sections[0])
            # The cover page is counted but not numbered when the rulebook
            # says so — it is page i, it just doesn't print the numeral.
            if not rulebook.page_numbering(rb).get(
                    "front_matter", {}).get("show_on_cover_page", True):
                _suppress_first_page_footer(doc.sections[0])

        # The body section (Chapter One onward) ALWAYS gets decimal numbering
        # restarting at 1 — this must happen even when there's no preliminary
        # content at all (e.g. a document that starts directly at Chapter
        # One with nothing before it), which previously left the whole
        # document with no page numbering configured at all.
        set_page_numbering(final_sectPr, "decimal", 1)
        _ensure_page_number_footer(doc.sections[-1])


    # Images sit where the rulebook says they sit.
    _image_align = _ALIGNMENT_BY_NAME.get(
        str(rulebook.figure_rules(rb).get("image_alignment") or "").lower())
    if _image_align is not None:
        for _p_el in doc.element.body.iter(qn("w:p")):
            if _p_el.find(".//" + qn("w:drawing")) is not None \
                    or _p_el.find(".//" + qn("w:pict")) is not None:
                Paragraph(_p_el, doc).alignment = _image_align

    # Caption placement, last of the structural passes: it moves
    # paragraphs, so nothing that depends on paragraph position may run
    # after it.
    reposition_captions(table_caption_paragraphs, figure_caption_paragraphs, rb)

    if directives.get("italicizeEtAl"):
        apply_italicize_phrase(doc, "et al")

    # Make Word populate the Table of Contents field (and page-number
    # fields) as soon as the student opens the file — see
    # _enable_update_fields_on_open.
    _enable_update_fields_on_open(doc)

    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()


# ---------------------------------------------------------------------------
# Field-code / footer helpers (raw OOXML — python-docx has no high-level API
# for these)
# ---------------------------------------------------------------------------
def _insert_toc_field(paragraph):
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose Update Field to generate the Table of Contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_separate)
    r_element.append(placeholder)
    r_element.append(fld_end)


def _enable_update_fields_on_open(doc):
    """Asks Word to refresh all fields when the document is opened.

    A native TOC field is empty until something refreshes it — without
    this the student opens the file and sees a placeholder line instead
    of their table of contents, which is indistinguishable from "the tool
    didn't generate one". Setting <w:updateFields w:val="true"/> makes
    Word populate the TOC (with real page numbers, which only a layout
    engine can know) the moment the file opens, so the generated table of
    contents is actually there for someone who just wants to print."""
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        settings.remove(existing)
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def _build_prelim_section_heading(doc, before_element, section_id, fallback_title=None,
                                  rb=None):
    """Creates a standalone front-matter heading paragraph immediately
    before `before_element`, styled from its rulebook section exactly like
    a heading the document already had. Position is provisional —
    PrelimAssembler repositions it — so this only needs a valid attachment
    point."""
    heading = doc.add_paragraph()
    before_element.addprevious(heading._p)
    _apply_prelim_heading(doc, heading, section_id,
                          fallback_title=fallback_title, rb=rb)
    return heading


def _build_list_section(doc, before_element, section_id, fallback_title, entries,
                        rb=None):
    """Builds a List of Tables/Figures/Plates section for a document that
    had none. Returns every paragraph it created, in reading order, for
    the caller to register with the assembler — returning them all is what
    keeps the heading and its entries together when the assembler later
    relocates the block.

    Goes through the same rulebook-driven builders the in-place path uses,
    so a generated list page and a rewritten one are indistinguishable.
    """
    if before_element is None:
        return []
    heading = _build_prelim_section_heading(doc, before_element, section_id,
                                            fallback_title, rb)
    return [heading] + _build_listing_rows(doc, heading, section_id, entries, rb)


def _build_toc_section(doc, before_element, rb=None):
    """Builds a Table of Contents heading plus a native Word TOC field.
    Returns both paragraphs for the caller to register."""
    if before_element is None:
        return []
    heading = _build_prelim_section_heading(
        doc, before_element, "table_of_contents", "TABLE OF CONTENTS", rb)
    field_p = insert_paragraph_after(heading, doc)
    _insert_toc_field(field_p)
    return [heading, field_p]


def _suppress_first_page_footer(section):
    """Turns on Word's "Different First Page" for this section and leaves
    that first-page footer empty.

    The rulebook counts the cover page as page i but does not print a
    numeral on it. Word has no way to say "count but do not show" other
    than giving the first page its own, blank, footer — which is what this
    does.
    """
    sectPr = section._sectPr
    if sectPr.find(qn("w:titlePg")) is None:
        title_pg = OxmlElement("w:titlePg")
        sectPr.append(title_pg)
    first = section.first_page_footer
    first.is_linked_to_previous = False
    for para in first.paragraphs:
        for run in list(para.runs):
            run._r.getparent().remove(run._r)


def _ensure_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs and footer.paragraphs[0].runs:
        return  # already has content, don't duplicate
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)


# ---------------------------------------------------------------------------
# Prelim page template generation (Cover/Title/Declaration/Certification)
# and AI-drafted content insertion (Dedication/Acknowledgement)
# ---------------------------------------------------------------------------
def _full_name(personal):
    parts = [personal.get("firstName", ""), personal.get("middleName", ""), personal.get("surname", "")]
    return " ".join(p.strip() for p in parts if p and p.strip()).upper()


# ---------------------------------------------------------------------------
# Full-page-height layout table (Cover Page / Title Page)
# ---------------------------------------------------------------------------
# A real Nigerian project cover/title page reads: title/BY/name/matric
# (and, on a Title Page, the submission sentence + supervisor) starting
# near the TOP of the page, then the submission month/year sitting near
# the BOTTOM margin — with nothing in between. A naive "just add
# paragraphs in order" approach leaves everything bunched at the top,
# since Word only pushes content down as far as the content itself
# requires; a "prepend N blank paragraphs" hack (the previous approach
# here) is fragile — too few and the date isn't pushed down enough, too
# many (combined with a long topic) and it overflows onto a second page.
#
# The reliable fix real templates use: a borderless, two-row, one-column
# table sized to span the section's full usable page height. The top row
# (vertically top-aligned, height "atLeast" so it grows for an unusually
# long title without clipping) holds the title/BY/name/matric block; the
# bottom row (vertically bottom-aligned, fixed height) holds the
# month/year line, landing right at the bottom margin regardless of how
# much text is above it, without overflowing for any normal-length title.
def _insert_in_tblPr_order(tblPr, new_el):
    for tag in _TBLPR_TAIL_TAGS:
        tail_el = tblPr.find(qn(tag))
        if tail_el is not None:
            tail_el.addprevious(new_el)
            return
    tblPr.append(new_el)


def _strip_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_border_edge(edge, "nil"))
    _insert_in_tblPr_order(tblPr, borders)


def _set_row_height(row, height, rule):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    existing = trPr.find(qn("w:trHeight"))
    if existing is not None:
        trPr.remove(existing)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(Emu(height).twips))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def _set_cell_vertical_alignment(cell, align):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _fill_cell_lines(cell, lines):
    """lines: list of (text, bold, size_pt) tuples. Reuses the cell's own
    default empty paragraph for the first line instead of leaving a
    stray blank line above the content."""
    first = True
    for text, bold, size_pt in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size_pt or BODY_PT)
        run.font.bold = bold


def build_full_height_layout_table(doc, anchor_element, top_lines, bottom_lines,
                                     bottom_row_height_in=1.3, page_break_before=False):
    """Inserts a borderless 2-row layout table immediately before
    anchor_element: top_lines top-aligned in the first row (grows if the
    content is unusually long), bottom_lines bottom-aligned in the second
    (fixed height, landing at the page's bottom margin). Together the two
    rows span the section's full usable page height."""
    section = doc.sections[0]
    usable_height = Emu(section.page_height - section.top_margin - section.bottom_margin)
    bottom_height = Inches(bottom_row_height_in)
    top_height = Emu(usable_height - bottom_height)
    if top_height < Inches(1):
        top_height = Emu(usable_height - Inches(0.8))
        bottom_height = Inches(0.8)

    table = doc.add_table(rows=2, cols=1)
    anchor_element.addprevious(table._tbl)
    table.autofit = False
    usable_width = section.page_width - section.left_margin - section.right_margin
    table.columns[0].width = usable_width
    for row in table.rows:
        row.cells[0].width = usable_width

    _strip_table_borders(table)

    top_cell, bottom_cell = table.cell(0, 0), table.cell(1, 0)
    _set_row_height(table.rows[0], top_height, "atLeast")
    _set_row_height(table.rows[1], bottom_height, "exact")
    _set_cell_vertical_alignment(top_cell, "top")
    _set_cell_vertical_alignment(bottom_cell, "bottom")

    _fill_cell_lines(top_cell, top_lines)
    _fill_cell_lines(bottom_cell, bottom_lines)

    if page_break_before:
        # A table can't carry page_break_before itself (that's a
        # paragraph-level property) — Word honors a break placed on the
        # FIRST paragraph inside the table's first cell instead.
        top_cell.paragraphs[0].paragraph_format.page_break_before = True

    return table


def _new_prepended_paragraph(doc, anchor_element):
    """Creates a new paragraph and places it immediately before
    anchor_element, which stays FIXED across repeated calls — each new
    paragraph lands right before the anchor, so calling this in the desired
    reading order naturally builds the block in correct order."""
    new_p = doc.add_paragraph()
    anchor_element.addprevious(new_p._p)
    return new_p


def _add_centered_line(doc, anchor, text, bold=False, size_pt=None, space_after_pt=10):
    p = _new_prepended_paragraph(doc, anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt or BODY_PT)
    run.font.bold = bold
    return p


def _add_signature_block(doc, anchor, name_line, role_label):
    """Standard Nigerian project signature block, matching the layout
    confirmed in real sample documents:
        {Name}                    ________________    ________________
        ({Role})                  Signature            Date
    """
    line1 = _new_prepended_paragraph(doc, anchor)
    line1.paragraph_format.tab_stops.add_tab_stop(Inches(3.5))
    line1.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
    r1 = line1.add_run(f"{name_line}\t________________\t________________")
    r1.font.name = FONT
    r1.font.size = Pt(BODY_PT)

    line2 = _new_prepended_paragraph(doc, anchor)
    line2.paragraph_format.tab_stops.add_tab_stop(Inches(3.5))
    line2.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
    r2 = line2.add_run(f"({role_label})\tSignature\tDate")
    r2.font.name = FONT
    r2.font.size = Pt(max(BODY_PT - 1, 8))
    line2.paragraph_format.space_after = Pt(10)
    return [line1, line2]


def _add_blank_examiner_line(doc, anchor):
    """Fallback when no External Examiner name is supplied: a blank
    underscore line with just the role label beneath — per explicit
    requirement, no invented name or signature/date columns."""
    line1 = _new_prepended_paragraph(doc, anchor)
    r1 = line1.add_run("________________")
    r1.font.name = FONT
    r1.font.size = Pt(BODY_PT)

    line2 = _new_prepended_paragraph(doc, anchor)
    r2 = line2.add_run("External Examiner")
    r2.font.name = FONT
    r2.font.size = Pt(max(BODY_PT - 1, 8))
    line2.paragraph_format.space_after = Pt(10)
    return [line1, line2]


def _case_text(text, case):
    """Applies a rulebook `case` value to one string."""
    text = text or ""
    case = str(case or "").lower()
    if case == "upper":
        return text.upper()
    if case == "lower":
        return text.lower()
    if case == "title":
        return text.title()
    return text


def _render_full_page_section(doc, anchor, section_id, values, rb=None,
                              page_break_before=False):
    """Builds a cover page or title page from its rulebook `elements` list.

    The rulebook decides which lines appear, in what order, in what case,
    how many blank paragraphs separate them, and which line is pinned to
    the bottom of the page. Nothing about the page's content is decided
    here — an element type this function has no value for is skipped, and
    a rulebook that drops a line (the cover page carries no supervisor,
    for instance) simply stops producing it.

    Returns the layout table, or None when the section has no elements.
    """
    section = rulebook.section_by_id(section_id, rb) or {}
    elements = section.get("elements") or []
    if not elements:
        return None

    top_lines = []
    bottom_lines = []

    for element in elements:
        etype = element.get("type")
        target = bottom_lines if element.get("anchor") == "bottom" else top_lines

        if etype == "spacer":
            for _ in range(int(element.get("paragraphs") or 1)):
                target.append(("", False, None))
            continue

        if element.get("template"):
            text = rulebook.fill_template(element["template"], values)
        else:
            text = values.get(etype)

        text = (text or "").strip()
        if not text:
            continue  # nothing supplied for this line — omit it, never invent

        size_pt = None
        if etype == "project_title":
            size_pt = rulebook.heading_size_pt("chapter_heading", BODY_PT, rb)
        target.append((_case_text(text, element.get("case")),
                       bool(element.get("bold", True)), size_pt))

    if not top_lines and not bottom_lines:
        return None

    return build_full_height_layout_table(
        doc, anchor,
        top_lines=top_lines,
        bottom_lines=bottom_lines,
        page_break_before=page_break_before,
    )


def _add_templated_body(doc, anchor, section_id, values, rb=None):
    """Writes a Declaration or Certification statement from the rulebook's
    own template, bolding the fields the rulebook marks."""
    spec = rulebook.section_body_spec(section_id, rb) or {}
    # A field can read differently inside a sentence than it does on the
    # cover page — the title is set in caps there and in title case here —
    # so the rulebook states the case per field.
    field_case = spec.get("field_case") or {}
    values = dict(values)
    for field, case in field_case.items():
        if values.get(field):
            values[field] = _case_text(values[field], case)

    text = rulebook.fill_template(spec.get("template"), values)
    if not text:
        return None

    para = _new_prepended_paragraph(doc, anchor)
    para.alignment = _ALIGNMENT_BY_NAME.get(
        str(spec.get("alignment") or "justify").lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)
    para.paragraph_format.line_spacing = rulebook.line_spacing_value(
        spec.get("line_spacing"), default=2.0)
    para.paragraph_format.space_after = Pt(20)

    # Split the finished sentence around the values the rulebook wants in
    # bold, so the emphasis lands on the student's own details rather than
    # on a fixed slice of the template.
    emphasis = []
    for field in spec.get("bold_fields") or []:
        value = (values.get(field) or "").strip()
        if value and value in text:
            emphasis.append(value)

    remaining = text
    while emphasis and remaining:
        hits = [(remaining.find(v), v) for v in emphasis if v in remaining]
        if not hits:
            break
        pos, value = min(hits)
        if pos > 0:
            run = para.add_run(remaining[:pos])
            run.font.name = FONT
            run.font.size = Pt(BODY_PT)
        run = para.add_run(value)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        run.font.bold = True
        remaining = remaining[pos + len(value):]
    if remaining:
        run = para.add_run(remaining)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
    return para


def _add_signature_rows(doc, anchor, section_id, values, rb=None):
    """Builds a section's signature block exactly as the rulebook lays it
    out: which rows, whose name on each, the label under each name, the
    tab positions of the signature and date lines, and the blank
    paragraphs between rows.

    A row whose name field is empty is dropped unless the rulebook marks
    it allow_blank_name — the External Supervisor line is printed blank on
    purpose, but a missing Dean is simply not printed rather than shown as
    an empty slot.
    """
    spec = rulebook.signature_block_spec(section_id, rb) or {}
    rows = spec.get("rows") or []
    if not rows:
        return []

    line_spacing = rulebook.line_spacing_value(spec.get("line_spacing"), default=1.0)
    space_before = Pt(spec.get("space_before_pt") or 0)
    space_after = Pt(spec.get("space_after_pt") or 0)
    tab_stops = spec.get("tab_stops_inches") or []
    label_pt = spec.get("label_size_pt") or BODY_PT
    gap = int(spec.get("spacer_paragraphs_between_rows") or 0)
    name_bold = bool(spec.get("name_bold"))
    sig_rule = "_" * int(spec.get("signature_rule_chars") or 20)
    date_rule = "_" * int(spec.get("date_rule_chars") or 14)
    created = []

    def new_line():
        para = _new_prepended_paragraph(doc, anchor)
        para.paragraph_format.line_spacing = line_spacing
        para.paragraph_format.space_before = space_before
        para.paragraph_format.space_after = space_after
        for stop in tab_stops:
            try:
                para.paragraph_format.tab_stops.add_tab_stop(Inches(float(stop)))
            except Exception:
                pass
        created.append(para)
        return para

    for position, row in enumerate(rows):
        name_field = row.get("name_field")
        name = (values.get(name_field) or "").strip() if name_field else ""
        if name_field and not name and not row.get("allow_blank_name"):
            continue
        if not name and row.get("blank_name_text"):
            # No name supplied, and the rulebook says what to print in its
            # place — the External Supervisor line is signed on the day, so
            # it carries its role instead of a name.
            name = rulebook.fill_template(row["blank_name_text"], values)

        # The rulebook names each column of the row, and the labels line
        # beneath it has one entry per column. Nothing is appended here
        # that the rulebook did not ask for, so a Declaration row that is
        # just a signature and a date does not sprout a third label.
        cells = []
        for column in row.get("columns") or []:
            if column == "name":
                cells.append(name)
            elif column == "signature_line":
                cells.append(sig_rule)
            elif column == "date_line":
                cells.append(date_rule)
            else:
                cells.append(rulebook.fill_template(column, values))

        top = new_line()
        run = top.add_run("\t".join(cells))
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        run.font.bold = name_bold and bool(name)

        labels = [rulebook.fill_template(label, values)
                  for label in (row.get("labels") or [])]
        if labels:
            bottom = new_line()
            run = bottom.add_run("\t".join(labels))
            run.font.name = FONT
            run.font.size = Pt(label_pt)

        if gap and position < len(rows) - 1:
            for _ in range(gap):
                created.append(_new_prepended_paragraph(doc, anchor))

    return created


def generate_prelim_pages(doc, prelim_toggles, personal, ai_content,
                          has_dedication, has_acknowledgement,
                          prelim=None, before_element=None):
    """Builds the front-matter sections the user asked to have generated —
    a replacement Cover/Title page, a template Declaration/Certification,
    or AI-drafted Dedication/Acknowledgement text for sections the original
    document genuinely lacked.

    Every block it creates is registered with `prelim` (the
    PrelimAssembler) at its canonical rank. Placement is NOT this
    function's job: it builds blocks in front of `before_element` purely
    to have somewhere valid to attach them, and the assembler moves them
    into canonical position afterwards. That separation is deliberate —
    when this function owned placement too, it anchored everything at the
    very top of the document, which is how an AI-drafted Acknowledgements
    page ended up printed ahead of a student's own cover page.
    """
    if before_element is None:
        if not doc.paragraphs:
            return
        before_element = doc.paragraphs[0]._p
    anchor = before_element

    def emit(rank, *objs):
        """Registers created blocks at their canonical rank. A no-op when
        called without an assembler, so this stays usable standalone."""
        if prelim is None:
            return
        for obj in objs:
            if isinstance(obj, (list, tuple)):
                prelim.add_all(rank, obj)
            else:
                prelim.add(rank, obj)

    full_name = _full_name(personal)
    topic = (personal.get("projectTopic") or "").strip().upper()
    matric = personal.get("matricNumber", "")
    month = personal.get("submissionMonth", "")
    year = personal.get("submissionYear", "")
    department = personal.get("department", "")
    faculty = personal.get("faculty", "")
    university = personal.get("university", "")
    degree = personal.get("degreeAwarded", "")
    supervisor = personal.get("supervisorName", "")
    hod = personal.get("hodName", "")
    dean = personal.get("deanName", "")
    examiner = (personal.get("externalExaminerName") or "").strip()
    date_line = (month.upper() + ", " + str(year)).strip(", ")

    # Every value the rulebook's templates can ask for, in one place. The
    # rulebook decides which of them each page actually uses.
    rb = rulebook.load_rulebook()
    template_values = {
        "project_title": topic,
        "student_name": full_name,
        "mat_number": matric,
        "month_year": date_line,
        "department": department,
        "faculty": faculty,
        "university": university,
        "degree": degree,
        "course_of_study": personal.get("courseOfStudy") or department,
        "supervisor_name": supervisor,
        "hod_name": hod,
        "dean_name": dean,
        "external_examiner_name": examiner,
    }

    # --- Cover Page ------------------------------------------------------
    # A full-page-height layout table (see build_full_height_layout_table)
    # so the lines spread from the top margin down to the date at the
    # bottom margin instead of bunching at the top, and so the page cannot
    # overflow onto a second sheet.
    if prelim_toggles.get("coverPage"):
        cover = _render_full_page_section(
            doc, anchor, "cover_page", template_values, rb)
        if cover is not None:
            emit(PRELIM_RANK["front"], cover)

    # --- Title Page ------------------------------------------------------
    if prelim_toggles.get("titlePage"):
        title_page = _render_full_page_section(
            doc, anchor, "title_page", template_values, rb,
            page_break_before=bool(prelim_toggles.get("coverPage")))
        if title_page is not None:
            emit(PRELIM_RANK["front"], title_page)

    # --- Declaration -----------------------------------------------------
    if prelim_toggles.get("declaration"):
        heading = _new_prepended_paragraph(doc, anchor)
        _apply_prelim_heading(doc, heading, "declaration",
                              fallback_title="DECLARATION", rb=rb)
        blocks = [heading]
        body = _add_templated_body(doc, anchor, "declaration", template_values, rb)
        if body is not None:
            blocks.append(body)
        decl_section = rulebook.section_by_id("declaration", rb) or {}
        for _ in range(int(decl_section.get("spacer_before_signatures") or 0)):
            blocks.append(_new_prepended_paragraph(doc, anchor))
        blocks.extend(_add_signature_rows(doc, anchor, "declaration",
                                          template_values, rb))
        emit(PRELIM_RANK["declaration"], blocks)

    # --- Certification ---------------------------------------------------
    if prelim_toggles.get("certification"):
        heading = _new_prepended_paragraph(doc, anchor)
        _apply_prelim_heading(doc, heading, "certification",
                              fallback_title="CERTIFICATION", rb=rb)
        blocks = [heading]
        body = _add_templated_body(doc, anchor, "certification", template_values, rb)
        if body is not None:
            blocks.append(body)
        cert_section = rulebook.section_by_id("certification", rb) or {}
        for _ in range(int(cert_section.get("spacer_before_signatures") or 0)):
            blocks.append(_new_prepended_paragraph(doc, anchor))
        blocks.extend(_add_signature_rows(doc, anchor, "certification",
                                          template_values, rb))
        emit(PRELIM_RANK["certification"], blocks)

    # --- Dedication (AI-drafted, only when genuinely missing) ---------------
    if prelim_toggles.get("dedication") and not has_dedication and ai_content.get("dedicationText"):
        heading = _add_centered_line(doc, anchor, "DEDICATION", bold=True, size_pt=CHAPTER_HEADING_PT, space_after_pt=15)
        heading.paragraph_format.page_break_before = True
        apply_heading_style(doc, heading, 1)
        body = _new_prepended_paragraph(doc, anchor)
        body.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body.paragraph_format.line_spacing = 1.5
        run = body.add_run(ai_content["dedicationText"].strip())
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        run.italic = True  # visually flags AI-drafted prose for the student
        # to review/edit before submitting — the UI also tells them this
        # explicitly, this is just a visual reinforcement in the file itself.
        emit(PRELIM_RANK["dedication"], heading, body)

    # --- Acknowledgement (AI-drafted, only when genuinely missing) ---------
    if prelim_toggles.get("acknowledgement") and not has_acknowledgement and ai_content.get("acknowledgementText"):
        heading = _add_centered_line(doc, anchor, "ACKNOWLEDGEMENTS", bold=True, size_pt=CHAPTER_HEADING_PT, space_after_pt=15)
        heading.paragraph_format.page_break_before = True
        apply_heading_style(doc, heading, 1)
        emit(PRELIM_RANK["acknowledgement"], heading)
        for para_text in ai_content["acknowledgementText"].strip().split("\n\n"):
            if not para_text.strip():
                continue
            body = _new_prepended_paragraph(doc, anchor)
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body.paragraph_format.line_spacing = 1.5
            run = body.add_run(para_text.strip())
            run.font.name = FONT
            run.font.size = Pt(BODY_PT)
            run.italic = True
            emit(PRELIM_RANK["acknowledgement"], body)


# ---------------------------------------------------------------------------
# Custom-prompt directive: italicize "et al." (or any other targeted phrase)
# ---------------------------------------------------------------------------
def apply_italicize_phrase(doc, phrase="et al"):
    """Splits any run containing the target phrase into up to three runs
    (before / phrase / after), italicizing only the middle one — a run's
    formatting applies to its entire text, so partial italics require
    splitting the run first."""
    pattern = re.compile(re.escape(phrase) + r"\.?", re.IGNORECASE)

    for p in doc.paragraphs:
        processed_run_ids = set()
        while True:
            match_found = False
            for run in list(p.runs):
                if id(run._r) in processed_run_ids:
                    continue
                text = run.text
                match = pattern.search(text)
                if not match:
                    continue
                match_found = True
                before, matched, after = text[: match.start()], text[match.start() : match.end()], text[match.end() :]

                # Capture ORIGINAL formatting before mutating — copying it
                # after setting italic=True would leak italic onto the
                # before/after split runs (a real bug caught in testing:
                # entire paragraphs were turning italic instead of just
                # "et al.").
                original_rPr = run._r.find(qn("w:rPr"))
                rpr_for_before = deepcopy(original_rPr) if original_rPr is not None else None
                rpr_for_after = deepcopy(original_rPr) if original_rPr is not None else None

                run.text = matched
                run.italic = True
                processed_run_ids.add(id(run._r))

                if before:
                    before_run = OxmlElement("w:r")
                    if rpr_for_before is not None:
                        before_run.append(rpr_for_before)
                    t_el = OxmlElement("w:t")
                    t_el.set(qn("xml:space"), "preserve")
                    t_el.text = before
                    before_run.append(t_el)
                    run._r.addprevious(before_run)

                if after:
                    after_run = OxmlElement("w:r")
                    if rpr_for_after is not None:
                        after_run.append(rpr_for_after)
                    t_el = OxmlElement("w:t")
                    t_el.set(qn("xml:space"), "preserve")
                    t_el.text = after
                    after_run.append(t_el)
                    run._r.addnext(after_run)

                break  # restart the scan from a fresh p.runs — the "after"
                # run just created might itself contain another match.

            if not match_found:
                break
